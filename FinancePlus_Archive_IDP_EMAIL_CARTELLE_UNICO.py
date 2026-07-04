#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
FinancePlus Archive AI IDP PRO COMPLETO
======================================
Web app Streamlit in un unico file .py per gestione intelligente documenti:
OCR, IDP, classificazione, riconoscimento cliente, rinomina, archivio,
SQLite, full-text search, coda da verificare, apprendimento, report, export,
backup, utenti, ricerca email azienda IMAP e predisposizione cloud/API.

Avvio:
    streamlit run FinancePlus_Archive_IDP_EMAIL_CARTELLE_UNICO.py

Dipendenze consigliate:
    pip install streamlit pandas pymupdf pillow pytesseract reportlab openpyxl python-docx python-dotenv

Nota sicurezza:
    Non inserire chiavi API nel codice. Usare .env locale o variabili ambiente.
"""
from __future__ import annotations

import csv
import email
import imaplib
import hashlib
import io
import json
import os
import re
import shutil
import sqlite3
import sys
import tempfile
import time
import traceback
import zipfile
from dataclasses import dataclass, asdict
from datetime import datetime, date
from email import policy
from email.header import decode_header, make_header
from email.message import EmailMessage
from email.utils import parsedate_to_datetime
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any, Optional

try:
    import streamlit as st
except Exception:
    st = None

try:
    import pandas as pd
except Exception:
    pd = None

try:
    from dotenv import load_dotenv
except Exception:
    load_dotenv = None

if load_dotenv:
    load_dotenv()

APP_TITLE = "FinancePlus Archive AI - IDP PRO Completo"
APP_VERSION = "3.3-idp-email-cartelle-unico-file"
ROOT_DIR = Path(__file__).resolve().parent
DATA_DIR = ROOT_DIR / "financeplus_archive_ai"
ARCHIVE_DIR = DATA_DIR / "Archivio_Documenti"
EMAIL_ARCHIVE_DIR = DATA_DIR / "Archivio_Email_Aziende"
VERIFY_DIR = ARCHIVE_DIR / "Da_Verificare"
TEMP_DIR = DATA_DIR / "temp"
BACKUP_DIR = DATA_DIR / "backup"
REPORT_DIR = DATA_DIR / "report"
LOG_DIR = DATA_DIR / "log"
DB_PATH = DATA_DIR / "financeplus_archive_ai.db"

SUPPORTED_EXTENSIONS = {
    ".pdf", ".txt", ".csv", ".xml", ".json",
    ".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp", ".webp",
    ".docx", ".xlsx", ".xlsm"
}
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp", ".webp"}
TEXT_EXTENSIONS = {".txt", ".csv", ".xml", ".json"}

MONTHS_IT = {
    1: "Gennaio", 2: "Febbraio", 3: "Marzo", 4: "Aprile", 5: "Maggio", 6: "Giugno",
    7: "Luglio", 8: "Agosto", 9: "Settembre", 10: "Ottobre", 11: "Novembre", 12: "Dicembre"
}

DOCUMENT_CATEGORIES = [
    "Visura camerale",
    "Bilancio",
    "Centrale Rischi",
    "DURC",
    "Documento di identita",
    "Contratto",
    "Fattura",
    "Estratto conto",
    "Preventivo",
    "Business plan",
    "Report bancario",
    "Dichiarazione fiscale",
    "Documento generico",
]

CATEGORY_FOLDERS = {
    "Visura camerale": "Visura",
    "Bilancio": "Bilancio",
    "Centrale Rischi": "Centrale_Rischi",
    "DURC": "DURC",
    "Documento di identita": "Documenti_Identita",
    "Contratto": "Contratti",
    "Fattura": "Fatture",
    "Estratto conto": "Estratti_Conto",
    "Preventivo": "Preventivi",
    "Business plan": "Business_Plan",
    "Report bancario": "Report_Bancari",
    "Dichiarazione fiscale": "Dichiarazioni_Fiscali",
    "Documento generico": "Altro",
}

CATEGORY_RULES = {
    "Visura camerale": [
        "visura", "registro imprese", "camera di commercio", "rea", "cciaa", "ateco",
        "capitale sociale", "sede legale", "amministratore unico", "codice attivita"
    ],
    "Bilancio": [
        "bilancio", "stato patrimoniale", "conto economico", "nota integrativa", "ricavi",
        "utile", "perdita", "patrimonio netto", "debiti", "ebitda", "ammortamenti"
    ],
    "Centrale Rischi": [
        "centrale rischi", "banca d'italia", "accordato", "utilizzato", "sconfini",
        "garanzie", "revoca", "autoliquidanti", "scaduto", "sofferenza"
    ],
    "DURC": [
        "durc", "regolarita contributiva", "inps", "inail", "cassa edile", "esito regolare",
        "numero protocollo", "validita"
    ],
    "Documento di identita": [
        "carta d'identita", "carta di identita", "passaporto", "patente", "documento identita",
        "comune di", "rilasciata", "scadenza", "nato a", "nata a"
    ],
    "Contratto": [
        "contratto", "scrittura privata", "accordo", "parti", "decorrenza", "durata",
        "clausola", "recesso", "penale", "mandato", "fornitura"
    ],
    "Fattura": [
        "fattura", "numero fattura", "imponibile", "iva", "totale documento", "codice destinatario",
        "sdi", "split payment", "fattura elettronica", "cedente", "cessionario"
    ],
    "Estratto conto": [
        "estratto conto", "lista movimenti", "conto corrente", "iban", "saldo iniziale",
        "saldo finale", "data valuta", "data contabile", "banca"
    ],
    "Preventivo": [
        "preventivo", "offerta", "validita offerta", "proposta economica", "fornitura",
        "prezzo", "totale preventivo", "condizioni di pagamento"
    ],
    "Business plan": [
        "business plan", "piano economico", "piano finanziario", "forecast", "previsionale",
        "break even", "investimento", "fabbisogno", "dscr", "cash flow"
    ],
    "Report bancario": [
        "report bancario", "dossier banca", "merito creditizio", "rating", "mcc", "garanzia",
        "durata finanziamento", "importo richiesto", "scoring", "analisi centrale rischi"
    ],
    "Dichiarazione fiscale": [
        "dichiarazione", "modello redditi", "irap", "iva", "f24", "quadro", "agenzia delle entrate",
        "protocollo telematico", "reddito imponibile", "imposte"
    ],
    "Documento generico": [],
}

ROLE_PERMS = {
    "Amministratore": {"read", "import", "review", "delete", "export", "backup", "users", "settings"},
    "Operatore": {"read", "import", "review", "export"},
    "Solo lettura": {"read"},
}

PIVA_RE = re.compile(r"(?:p\.?\s*iva|partita\s+iva|piva|vat)\s*[:\-.]?\s*([0-9]{11})", re.I)
CF_RE = re.compile(r"(?:codice\s+fiscale|c\.?\s*f\.?)\s*[:\-.]?\s*([A-Z0-9]{11,16})\b", re.I)
CF_PERSON_RE = re.compile(r"\b([A-Z]{6}[0-9]{2}[A-Z][0-9]{2}[A-Z][0-9]{3}[A-Z])\b", re.I)
EMAIL_RE = re.compile(r"[A-Z0-9._%+\-]+@[A-Z0-9.\-]+\.[A-Z]{2,}", re.I)
PEC_RE = re.compile(r"[A-Z0-9._%+\-]+@[A-Z0-9.\-]*(?:pec|legalmail|arubapec|postacert)[A-Z0-9.\-]*\.[A-Z]{2,}", re.I)
ATECO_RE = re.compile(r"(?:ateco|codice\s+attivita)\s*[:\-.]?\s*([0-9]{2}(?:\.[0-9]{1,2}){0,3}|[0-9]{6})", re.I)
IBAN_RE = re.compile(r"\bIT[0-9]{2}[A-Z][0-9A-Z]{22}\b", re.I)
DATE_RE = re.compile(r"\b([0-3]?\d[\-/\.][01]?\d[\-/\.](?:19|20)?\d{2}|(?:19|20)\d{2}[\-/\.][01]?\d[\-/\.][0-3]?\d)\b")
MONEY_RE = re.compile(r"(?:EUR|euro|€)?\s*([0-9]{1,3}(?:\.[0-9]{3})*(?:,[0-9]{2})|[0-9]+(?:,[0-9]{2})?)\s*(?:EUR|euro|€)?", re.I)
LEGAL_SUFFIX = r"(?:S\.?\s*R\.?\s*L\.?|SRL|SOCIETA'?\s+A\s+RESPONSABILITA'?\s+LIMITATA|S\.?\s*P\.?\s*A\.?|SPA|S\.?\s*A\.?\s*S\.?|SAS|S\.?\s*N\.?\s*C\.?|SNC|SOC\.?\s*COOP\.?|COOPERATIVA)"

@dataclass
class ExtractionResult:
    text: str
    engine: str
    confidence: float
    pages: int = 0
    warnings: list[str] | None = None

@dataclass
class ClassificationResult:
    categoria: str
    confidence: float
    ragione_sociale: str = ""
    partita_iva: str = ""
    codice_fiscale: str = ""
    sede: str = ""
    amministratore: str = ""
    ateco: str = ""
    email: str = ""
    pec: str = ""
    data_documento: str = ""
    importo: float | None = None
    iban: str = ""
    protocollo: str = ""
    metadata: dict[str, Any] | None = None
    motivazione: str = ""
    anomalie: list[str] | None = None
    needs_review: bool = True
    matched_client_id: int | None = None


def now_iso() -> str:
    return datetime.now().isoformat(timespec="seconds")


def ensure_dirs() -> None:
    for p in [DATA_DIR, ARCHIVE_DIR, VERIFY_DIR, TEMP_DIR, BACKUP_DIR, REPORT_DIR, LOG_DIR]:
        p.mkdir(parents=True, exist_ok=True)


def normalize_spaces(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def strip_accents_basic(value: str) -> str:
    table = str.maketrans({
        "à": "a", "è": "e", "é": "e", "ì": "i", "ò": "o", "ù": "u",
        "À": "A", "È": "E", "É": "E", "Ì": "I", "Ò": "O", "Ù": "U",
        "’": "'", "`": "'",
    })
    return value.translate(table)


def safe_filename(value: str, max_len: int = 90, default: str = "SENZA_NOME") -> str:
    value = strip_accents_basic(normalize_spaces(value or default))
    value = re.sub(r"[\\/:*?\"<>|]+", "_", value)
    value = re.sub(r"[^A-Za-z0-9_ .\-'&]+", "_", value)
    value = re.sub(r"\s+", "_", value)
    value = re.sub(r"_+", "_", value).strip(" ._-")
    if not value:
        value = default
    if len(value) > max_len:
        h = hashlib.sha1(value.encode("utf-8", errors="ignore")).hexdigest()[:8]
        value = value[: max_len - 9].rstrip(" ._-") + "_" + h
    return value


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def hash_password(password: str) -> str:
    return hashlib.sha256(("financeplus_idp::" + password).encode("utf-8")).hexdigest()


def connect() -> sqlite3.Connection:
    ensure_dirs()
    conn = sqlite3.connect(str(DB_PATH), check_same_thread=False)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys=ON")
    return conn


def init_db() -> None:
    ensure_dirs()
    con = connect()
    try:
        con.executescript(
            """
            PRAGMA journal_mode=WAL;
            PRAGMA foreign_keys=ON;

            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE NOT NULL,
                password_hash TEXT NOT NULL,
                role TEXT NOT NULL DEFAULT 'Operatore',
                active INTEGER NOT NULL DEFAULT 1,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS clienti (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                ragione_sociale TEXT NOT NULL,
                partita_iva TEXT,
                codice_fiscale TEXT,
                sede TEXT,
                amministratore TEXT,
                ateco TEXT,
                email TEXT,
                pec TEXT,
                data_creazione TEXT NOT NULL,
                data_ultimo_aggiornamento TEXT NOT NULL
            );

            CREATE UNIQUE INDEX IF NOT EXISTS idx_clienti_piva ON clienti(partita_iva) WHERE partita_iva IS NOT NULL AND partita_iva <> '';
            CREATE INDEX IF NOT EXISTS idx_clienti_cf ON clienti(codice_fiscale);
            CREATE INDEX IF NOT EXISTS idx_clienti_ragione ON clienti(ragione_sociale);

            CREATE TABLE IF NOT EXISTS documenti (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                cliente_id INTEGER,
                categoria TEXT NOT NULL,
                nome_file_originale TEXT NOT NULL,
                nome_file_archiviato TEXT NOT NULL,
                percorso_file TEXT NOT NULL,
                data_documento TEXT,
                importo REAL,
                hash_sha256 TEXT UNIQUE NOT NULL,
                testo_estratto TEXT,
                stato_riconoscimento TEXT NOT NULL,
                confidenza REAL NOT NULL DEFAULT 0,
                motore_estrazione TEXT,
                metadata_json TEXT,
                data_importazione TEXT NOT NULL,
                data_ultimo_aggiornamento TEXT NOT NULL,
                FOREIGN KEY(cliente_id) REFERENCES clienti(id)
            );

            CREATE INDEX IF NOT EXISTS idx_documenti_hash ON documenti(hash_sha256);
            CREATE INDEX IF NOT EXISTS idx_documenti_categoria ON documenti(categoria);
            CREATE INDEX IF NOT EXISTS idx_documenti_cliente ON documenti(cliente_id);
            CREATE INDEX IF NOT EXISTS idx_documenti_stato ON documenti(stato_riconoscimento);

            CREATE TABLE IF NOT EXISTS coda_da_verificare (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                documento_id INTEGER,
                file TEXT,
                testo_estratto TEXT,
                categoria_suggerita TEXT,
                cliente_suggerito TEXT,
                motivo_incertezza TEXT,
                stato_lavorazione TEXT NOT NULL DEFAULT 'aperta',
                data_creazione TEXT NOT NULL,
                data_chiusura TEXT,
                FOREIGN KEY(documento_id) REFERENCES documenti(id)
            );

            CREATE TABLE IF NOT EXISTS log_attivita (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                data TEXT NOT NULL,
                utente TEXT,
                operazione TEXT NOT NULL,
                documento TEXT,
                risultato TEXT,
                note TEXT
            );

            CREATE TABLE IF NOT EXISTS modello_apprendimento (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                testo_parole_chiave TEXT,
                categoria_corretta TEXT,
                cliente_corretto TEXT,
                partita_iva TEXT,
                codice_fiscale TEXT,
                regola_appresa TEXT,
                data_correzione TEXT NOT NULL,
                fonte_correzione TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS email_archiviate (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                azienda TEXT NOT NULL,
                mittente TEXT,
                destinatari TEXT,
                data_email TEXT,
                oggetto TEXT,
                message_id TEXT,
                mailbox TEXT,
                pdf_path TEXT,
                allegati_json TEXT,
                hash_message TEXT UNIQUE,
                score_abbinamento REAL DEFAULT 0,
                motivo_abbinamento TEXT,
                data_archiviazione TEXT NOT NULL
            );

            CREATE INDEX IF NOT EXISTS idx_email_arch_azienda ON email_archiviate(azienda);
            CREATE INDEX IF NOT EXISTS idx_email_arch_mittente ON email_archiviate(mittente);
            CREATE INDEX IF NOT EXISTS idx_email_arch_data ON email_archiviate(data_email);

            CREATE TABLE IF NOT EXISTS email_file_hash (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                hash_sha256 TEXT UNIQUE NOT NULL,
                percorso_file TEXT NOT NULL,
                azienda TEXT,
                mittente TEXT,
                nome_file TEXT,
                fonte TEXT,
                data_archiviazione TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS settings (
                key TEXT PRIMARY KEY,
                value TEXT,
                updated_at TEXT NOT NULL
            );
            """
        )
        try:
            con.execute(
                "CREATE VIRTUAL TABLE IF NOT EXISTS documenti_fts USING fts5(documento_id UNINDEXED, titolo, corpo, metadata)"
            )
        except sqlite3.OperationalError:
            pass
        admin = con.execute("SELECT id FROM users WHERE username=?", ("admin",)).fetchone()
        if not admin:
            con.execute(
                "INSERT INTO users(username,password_hash,role,active,created_at,updated_at) VALUES(?,?,?,?,?,?)",
                ("admin", hash_password("admin123"), "Amministratore", 1, now_iso(), now_iso()),
            )
        con.commit()
    finally:
        con.close()


def query_all(sql: str, params: tuple = ()) -> list[dict[str, Any]]:
    con = connect()
    try:
        return [dict(r) for r in con.execute(sql, params).fetchall()]
    finally:
        con.close()


def query_one(sql: str, params: tuple = ()) -> Optional[dict[str, Any]]:
    con = connect()
    try:
        row = con.execute(sql, params).fetchone()
        return dict(row) if row else None
    finally:
        con.close()


def execute(sql: str, params: tuple = ()) -> int:
    con = connect()
    try:
        cur = con.execute(sql, params)
        con.commit()
        return int(cur.lastrowid or 0)
    finally:
        con.close()


def log_event(operazione: str, documento: str = "", risultato: str = "OK", note: str = "") -> None:
    user = ""
    try:
        if st is not None:
            user = (st.session_state.get("user") or {}).get("username", "")
    except Exception:
        user = ""
    try:
        execute(
            "INSERT INTO log_attivita(data,utente,operazione,documento,risultato,note) VALUES(?,?,?,?,?,?)",
            (now_iso(), user, operazione, documento, risultato, note[:3000]),
        )
    except Exception:
        pass


def set_setting(key: str, value: str) -> None:
    execute(
        "INSERT INTO settings(key,value,updated_at) VALUES(?,?,?) ON CONFLICT(key) DO UPDATE SET value=excluded.value, updated_at=excluded.updated_at",
        (key, value, now_iso()),
    )


def get_setting(key: str, default: str = "") -> str:
    row = query_one("SELECT value FROM settings WHERE key=?", (key,))
    return str(row["value"]) if row else default


def parse_date_any(value: str | None) -> Optional[datetime]:
    if not value:
        return None
    v = str(value).strip()
    formats = ["%d/%m/%Y", "%d-%m-%Y", "%d.%m.%Y", "%d/%m/%y", "%d-%m-%y", "%Y-%m-%d", "%Y/%m/%d"]
    for fmt in formats:
        try:
            d = datetime.strptime(v, fmt)
            if d.year < 100:
                d = d.replace(year=2000 + d.year)
            return d
        except Exception:
            pass
    return None


def ddmmyyyy(value: datetime | date | None) -> str:
    if not value:
        value = datetime.now()
    if isinstance(value, date) and not isinstance(value, datetime):
        value = datetime(value.year, value.month, value.day)
    return value.strftime("%d-%m-%Y")


def yyyy_mm_dd(value: datetime | date | None) -> str:
    if not value:
        return ""
    if isinstance(value, date) and not isinstance(value, datetime):
        value = datetime(value.year, value.month, value.day)
    return value.strftime("%Y-%m-%d")


def month_folder(value: datetime | date | None) -> tuple[str, str]:
    if not value:
        value = datetime.now()
    return str(value.year), MONTHS_IT.get(value.month, str(value.month))


def unique_path(path: Path) -> Path:
    if not path.exists():
        return path
    stem, suffix = path.stem, path.suffix
    for i in range(2, 10000):
        candidate = path.with_name(f"{stem}_{i}{suffix}")
        if not candidate.exists():
            return candidate
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return path.with_name(f"{stem}_{stamp}{suffix}")


def clean_text(text: str, limit: int = 120000) -> str:
    text = text or ""
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text[:limit]


def extract_text_from_pdf(path: Path, use_local_ocr: bool, max_pages_ocr: int) -> ExtractionResult:
    warnings: list[str] = []
    text_parts: list[str] = []
    pages = 0
    engine = "PyMuPDF"
    try:
        import fitz
        with fitz.open(str(path)) as doc:
            pages = len(doc)
            for idx, page in enumerate(doc):
                if idx >= 80:
                    warnings.append("PDF lungo: testo limitato alle prime 80 pagine.")
                    break
                page_text = page.get_text("text") or ""
                text_parts.append(f"\n--- Pagina {idx + 1} ---\n{page_text}")
            text = clean_text("\n".join(text_parts))
            if len(normalize_spaces(text)) >= 80:
                return ExtractionResult(text=text, engine=engine, confidence=0.92, pages=pages, warnings=warnings)
            if use_local_ocr:
                ocr_parts = []
                try:
                    from PIL import Image
                    import pytesseract
                    limit = min(pages, max_pages_ocr)
                    with fitz.open(str(path)) as doc2:
                        for idx in range(limit):
                            page = doc2[idx]
                            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                            img = Image.open(io.BytesIO(pix.tobytes("png")))
                            try:
                                page_ocr = pytesseract.image_to_string(img, lang="ita+eng")
                            except Exception:
                                page_ocr = pytesseract.image_to_string(img)
                            ocr_parts.append(f"\n--- OCR Pagina {idx + 1} ---\n{page_ocr}")
                    engine = "PyMuPDF + pytesseract OCR"
                    text = clean_text("\n".join(ocr_parts))
                    conf = 0.74 if len(normalize_spaces(text)) >= 80 else 0.18
                    if pages > max_pages_ocr:
                        warnings.append(f"OCR limitato alle prime {max_pages_ocr} pagine.")
                    return ExtractionResult(text=text, engine=engine, confidence=conf, pages=pages, warnings=warnings)
                except Exception as exc:
                    warnings.append(f"OCR PDF non disponibile: {exc}")
                    return ExtractionResult(text=text, engine=engine, confidence=0.12, pages=pages, warnings=warnings)
            return ExtractionResult(text=text, engine=engine, confidence=0.15, pages=pages, warnings=warnings)
    except Exception as exc:
        warnings.append(f"Errore PyMuPDF: {exc}")
    try:
        try:
            from pypdf import PdfReader
        except Exception:
            from PyPDF2 import PdfReader  # type: ignore
        reader = PdfReader(str(path))
        pages = len(reader.pages)
        parts = []
        for idx, page in enumerate(reader.pages[:80]):
            parts.append(f"\n--- Pagina {idx + 1} ---\n{page.extract_text() or ''}")
        text = clean_text("\n".join(parts))
        conf = 0.70 if len(normalize_spaces(text)) >= 80 else 0.12
        return ExtractionResult(text=text, engine="pypdf", confidence=conf, pages=pages, warnings=warnings)
    except Exception as exc:
        warnings.append(f"PDF non leggibile: {exc}")
        return ExtractionResult(text="", engine="errore_pdf", confidence=0.0, pages=pages, warnings=warnings)


def extract_text_from_image(path: Path) -> ExtractionResult:
    warnings: list[str] = []
    try:
        from PIL import Image
        import pytesseract
        img = Image.open(str(path))
        try:
            text = pytesseract.image_to_string(img, lang="ita+eng")
        except Exception:
            text = pytesseract.image_to_string(img)
        text = clean_text(text)
        conf = 0.72 if len(normalize_spaces(text)) >= 60 else 0.12
        return ExtractionResult(text=text, engine="pytesseract OCR immagine", confidence=conf, pages=1, warnings=warnings)
    except Exception as exc:
        warnings.append(f"OCR immagine non disponibile: {exc}")
        return ExtractionResult(text="", engine="errore_ocr_immagine", confidence=0.0, pages=1, warnings=warnings)


def extract_text_from_docx(path: Path) -> ExtractionResult:
    warnings: list[str] = []
    try:
        from docx import Document
        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                cells = [normalize_spaces(cell.text) for cell in row.cells if normalize_spaces(cell.text)]
                if cells:
                    parts.append(" | ".join(cells))
        text = clean_text("\n".join(parts))
        return ExtractionResult(text=text, engine="python-docx", confidence=0.85 if len(text) > 60 else 0.12)
    except Exception as exc:
        warnings.append(f"DOCX non leggibile: {exc}")
        return ExtractionResult(text="", engine="errore_docx", confidence=0.0, warnings=warnings)


def extract_text_from_xlsx(path: Path) -> ExtractionResult:
    warnings: list[str] = []
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
        parts = []
        for ws in wb.worksheets[:10]:
            parts.append(f"--- Foglio {ws.title} ---")
            row_count = 0
            for row in ws.iter_rows(max_row=250, values_only=True):
                vals = [str(v) for v in row if v not in (None, "")]
                if vals:
                    parts.append(" | ".join(vals))
                row_count += 1
                if row_count >= 250:
                    break
        wb.close()
        text = clean_text("\n".join(parts))
        return ExtractionResult(text=text, engine="openpyxl", confidence=0.82 if len(text) > 60 else 0.10)
    except Exception as exc:
        warnings.append(f"Excel non leggibile: {exc}")
        return ExtractionResult(text="", engine="errore_xlsx", confidence=0.0, warnings=warnings)


def extract_text(path: Path, use_local_ocr: bool = True, use_cloud_ocr: bool = False, max_pages_ocr: int = 8) -> ExtractionResult:
    ext = path.suffix.lower()
    if ext == ".pdf":
        result = extract_text_from_pdf(path, use_local_ocr=use_local_ocr, max_pages_ocr=max_pages_ocr)
    elif ext in IMAGE_EXTENSIONS:
        result = extract_text_from_image(path) if use_local_ocr else ExtractionResult("", "ocr_disattivato", 0.0, 1, ["OCR locale disattivato."])
    elif ext in TEXT_EXTENSIONS:
        warnings = []
        for enc in ["utf-8-sig", "utf-8", "cp1252", "latin-1"]:
            try:
                text = clean_text(path.read_text(encoding=enc, errors="ignore"))
                return ExtractionResult(text=text, engine=f"testo {enc}", confidence=0.90 if len(text) > 40 else 0.20, warnings=warnings)
            except Exception:
                continue
        result = ExtractionResult("", "errore_testo", 0.0, warnings=["File testuale non leggibile."])
    elif ext == ".docx":
        result = extract_text_from_docx(path)
    elif ext in {".xlsx", ".xlsm"}:
        result = extract_text_from_xlsx(path)
    else:
        result = ExtractionResult("", "estensione_non_supportata", 0.0, warnings=["Estensione non supportata."])
    if use_cloud_ocr and result.confidence < 0.50:
        w = result.warnings or []
        w.append("OCR cloud/HTR predisposto: nessun invio eseguito senza adapter e credenziali configurate.")
        result.warnings = w
    return result


def first_match(regex: re.Pattern[str], text: str) -> str:
    m = regex.search(text or "")
    return normalize_spaces(m.group(1)) if m else ""


def extract_piva(text: str) -> str:
    p = first_match(PIVA_RE, text)
    if p:
        return p
    for m in re.finditer(r"\b([0-9]{11})\b", text or ""):
        ctx = (text[max(0, m.start() - 120): m.end() + 80] or "").lower()
        if any(k in ctx for k in ["iva", "fiscale", "rea", "impresa", "societa", "fattura", "cliente"]):
            return m.group(1)
    return ""


def extract_cf(text: str) -> str:
    cf = first_match(CF_RE, text)
    if cf:
        return cf.upper()
    m = CF_PERSON_RE.search(text or "")
    return m.group(1).upper() if m else ""


def extract_email(text: str) -> str:
    m = EMAIL_RE.search(text or "")
    return m.group(0).lower() if m else ""


def extract_pec(text: str) -> str:
    m = PEC_RE.search(text or "")
    return m.group(0).lower() if m else ""


def extract_ateco(text: str) -> str:
    return first_match(ATECO_RE, text)


def extract_iban(text: str) -> str:
    m = IBAN_RE.search((text or "").replace(" ", ""))
    return m.group(0).upper() if m else ""


def extract_dates(text: str, limit: int = 20) -> list[datetime]:
    out = []
    for raw in DATE_RE.findall(text or "")[:limit]:
        d = parse_date_any(raw)
        if d and 1990 <= d.year <= 2100:
            out.append(d)
    return out


def extract_document_date(text: str) -> str:
    dates = extract_dates(text, limit=30)
    if not dates:
        return ""
    return yyyy_mm_dd(dates[0])


def parse_amount(value: str) -> Optional[float]:
    if not value:
        return None
    s = value.replace("EUR", "").replace("euro", "").replace("€", "").replace(" ", "")
    if "," in s and "." in s:
        s = s.replace(".", "").replace(",", ".")
    elif "," in s:
        s = s.replace(",", ".")
    try:
        return float(s)
    except Exception:
        return None


def extract_amounts(text: str) -> list[float]:
    amounts = []
    for m in MONEY_RE.finditer(text or ""):
        val = parse_amount(m.group(1))
        if val is not None and val >= 1:
            amounts.append(val)
    return amounts


def extract_primary_amount(text: str) -> Optional[float]:
    patterns = [
        r"(?:totale\s+(?:documento|fattura|preventivo)?|importo\s+(?:richiesto|totale)?|saldo\s+finale|totale\s+da\s+pagare)[^0-9]{0,50}([0-9]{1,3}(?:\.[0-9]{3})*(?:,[0-9]{2})|[0-9]+(?:,[0-9]{2})?)",
        r"€\s*([0-9]{1,3}(?:\.[0-9]{3})*(?:,[0-9]{2})|[0-9]+(?:,[0-9]{2})?)",
    ]
    found = []
    for p in patterns:
        for m in re.finditer(p, text or "", flags=re.I):
            val = parse_amount(m.group(1))
            if val is not None:
                found.append(val)
    if found:
        return max(found)
    all_amounts = extract_amounts(text)
    return max(all_amounts) if all_amounts else None


def extract_ragione_sociale(text: str, filename: str = "") -> str:
    sample = (text or "")[:30000]
    patterns = [
        rf"(?:denominazione|ragione\s+sociale|impresa|azienda|societa|cliente|spettabile)\s*[:\-.]?\s*([A-Z0-9][A-Z0-9 '&\.\-]{{2,100}}\s+{LEGAL_SUFFIX})",
        rf"\b([A-Z0-9][A-Z0-9 '&\.\-]{{2,100}}\s+{LEGAL_SUFFIX})\b",
        rf"(?:denominazione|ragione\s+sociale|impresa|azienda|societa|cliente|spettabile)\s*[:\-.]?\s*([A-Z0-9][A-Z0-9 '&\.\-]{{4,100}})",
    ]
    bad_prefixes = ("registro", "camera", "agenzia", "banca", "centrale", "documento", "fattura", "bilancio")
    for p in patterns:
        for m in re.finditer(p, sample, flags=re.I):
            raw = normalize_spaces(m.group(1)).strip(" -:.;,")
            raw = re.split(r"\b(partita|p\.?iva|codice|c\.?f\.?|rea|pec|sede|tel|email|ateco|capitale)\b", raw, flags=re.I)[0]
            raw = normalize_spaces(raw).strip(" -:.;,")
            if len(raw) >= 4 and not raw.lower().startswith(bad_prefixes):
                return raw.upper()
    if filename:
        stem = re.sub(r"[_\-]+", " ", Path(filename).stem)
        stem = re.sub(r"(?i)\b(visura|bilancio|centrale|rischi|estratto|conto|fattura|preventivo|documento|report|business|plan|durc|signed|firmato|copia|pdf)\b", " ", stem)
        stem = normalize_spaces(stem).strip(" -_.")
        if len(stem) >= 4:
            return stem.upper()[:90]
    return ""


def extract_labeled(text: str, labels: list[str], max_len: int = 120) -> str:
    label_part = "|".join(re.escape(x) for x in labels)
    p = re.compile(rf"(?:{label_part})\s*[:\-.]?\s*([^\n\r]{{3,{max_len}}})", re.I)
    m = p.search(text or "")
    return normalize_spaces(m.group(1)).strip(" -:.;,") if m else ""


def extract_sede(text: str) -> str:
    return extract_labeled(text, ["sede legale", "sede", "indirizzo"], 150).upper()


def extract_amministratore(text: str) -> str:
    return extract_labeled(text, ["amministratore unico", "amministratore", "legale rappresentante", "rappresentante legale", "presidente"], 100).upper()


def extract_protocollo(text: str) -> str:
    p = extract_labeled(text, ["protocollo", "numero protocollo", "prot"], 70)
    if p:
        return p
    m = re.search(r"\b(?:prot\.?|protocollo)\s*[:\-.]?\s*([A-Z0-9/\-\.]{4,40})\b", text or "", flags=re.I)
    return m.group(1) if m else ""


def keyword_scores(text: str, filename: str = "") -> dict[str, float]:
    blob = (filename + "\n" + (text or "")[:50000]).lower()
    scores = {}
    for category, keys in CATEGORY_RULES.items():
        if not keys:
            scores[category] = 0.0
            continue
        hits = 0.0
        for key in keys:
            if key.lower() in blob:
                hits += 1.0
        scores[category] = hits / max(1, len(keys))
    return scores


def top_keywords(text: str, limit: int = 18) -> list[str]:
    words = re.findall(r"[A-Za-z0-9]{4,}", strip_accents_basic((text or "").lower()))
    stop = {
        "della", "delle", "degli", "alla", "allo", "alle", "sono", "come", "euro", "data",
        "documento", "codice", "fiscale", "partita", "pagina", "totale", "numero", "societa",
        "cliente", "firma", "sede", "legale", "con", "per", "del", "dei", "nel", "nella"
    }
    freq: dict[str, int] = {}
    for w in words:
        if w in stop or len(w) < 4:
            continue
        freq[w] = freq.get(w, 0) + 1
    return [w for w, _ in sorted(freq.items(), key=lambda x: x[1], reverse=True)[:limit]]


def learned_hint(text: str, piva: str, cf: str, ragione: str) -> dict[str, Any]:
    rows = query_all("SELECT * FROM modello_apprendimento ORDER BY id DESC LIMIT 1000")
    blob = strip_accents_basic((text or "").lower())
    best: dict[str, Any] = {"score": 0.0}
    for r in rows:
        score = 0.0
        if piva and r.get("partita_iva") == piva:
            score += 0.65
        if cf and r.get("codice_fiscale") == cf:
            score += 0.55
        if ragione and r.get("cliente_corretto"):
            score += 0.30 * SequenceMatcher(None, ragione.lower(), str(r.get("cliente_corretto")).lower()).ratio()
        keys = [k.strip().lower() for k in str(r.get("testo_parole_chiave") or "").split(",") if k.strip()]
        for k in keys[:18]:
            if k and k in blob:
                score += 0.04
        if score > best["score"]:
            best = {"score": score, "row": r}
    return best if best["score"] >= 0.55 else {"score": 0.0}


def find_client(piva: str = "", cf: str = "", ragione: str = "", email: str = "", pec: str = "", amministratore: str = "") -> Optional[dict[str, Any]]:
    if piva:
        row = query_one("SELECT * FROM clienti WHERE partita_iva=?", (piva,))
        if row:
            return row
    if cf:
        row = query_one("SELECT * FROM clienti WHERE codice_fiscale=?", (cf,))
        if row:
            return row
    if pec:
        row = query_one("SELECT * FROM clienti WHERE pec=?", (pec,))
        if row:
            return row
    if email:
        row = query_one("SELECT * FROM clienti WHERE email=?", (email,))
        if row:
            return row
    rows = query_all("SELECT * FROM clienti")
    best = None
    best_score = 0.0
    for r in rows:
        score = 0.0
        if ragione and r.get("ragione_sociale"):
            score = max(score, SequenceMatcher(None, ragione.lower(), str(r["ragione_sociale"]).lower()).ratio())
        if amministratore and r.get("amministratore"):
            score = max(score, 0.75 * SequenceMatcher(None, amministratore.lower(), str(r["amministratore"]).lower()).ratio())
        if score > best_score:
            best_score = score
            best = r
    if best and best_score >= 0.84:
        return best
    return None


def ensure_client(c: ClassificationResult, create: bool = True) -> Optional[int]:
    existing = find_client(c.partita_iva, c.codice_fiscale, c.ragione_sociale, c.email, c.pec, c.amministratore)
    if existing:
        updates = {
            "ragione_sociale": c.ragione_sociale or existing.get("ragione_sociale") or "",
            "partita_iva": c.partita_iva or existing.get("partita_iva") or "",
            "codice_fiscale": c.codice_fiscale or existing.get("codice_fiscale") or "",
            "sede": c.sede or existing.get("sede") or "",
            "amministratore": c.amministratore or existing.get("amministratore") or "",
            "ateco": c.ateco or existing.get("ateco") or "",
            "email": c.email or existing.get("email") or "",
            "pec": c.pec or existing.get("pec") or "",
            "data_ultimo_aggiornamento": now_iso(),
            "id": existing["id"],
        }
        execute(
            """
            UPDATE clienti SET ragione_sociale=?, partita_iva=?, codice_fiscale=?, sede=?, amministratore=?,
                ateco=?, email=?, pec=?, data_ultimo_aggiornamento=? WHERE id=?
            """,
            (
                updates["ragione_sociale"], updates["partita_iva"], updates["codice_fiscale"], updates["sede"],
                updates["amministratore"], updates["ateco"], updates["email"], updates["pec"], updates["data_ultimo_aggiornamento"], updates["id"]
            ),
        )
        return int(existing["id"])
    if not create:
        return None
    if not (c.ragione_sociale or c.partita_iva or c.codice_fiscale):
        return None
    ragione = c.ragione_sociale or f"CLIENTE_{c.partita_iva or c.codice_fiscale or 'DA_VERIFICARE'}"
    try:
        return execute(
            """
            INSERT INTO clienti(ragione_sociale,partita_iva,codice_fiscale,sede,amministratore,ateco,email,pec,data_creazione,data_ultimo_aggiornamento)
            VALUES(?,?,?,?,?,?,?,?,?,?)
            """,
            (ragione, c.partita_iva, c.codice_fiscale, c.sede, c.amministratore, c.ateco, c.email, c.pec, now_iso(), now_iso()),
        )
    except sqlite3.IntegrityError:
        row = find_client(c.partita_iva, c.codice_fiscale, ragione, c.email, c.pec, c.amministratore)
        return int(row["id"]) if row else None


def extract_category_metadata(category: str, text: str) -> dict[str, Any]:
    t = text or ""
    meta: dict[str, Any] = {}
    if category == "Visura camerale":
        meta.update({
            "sede_legale": extract_sede(t),
            "amministratore": extract_amministratore(t),
            "codice_ateco": extract_ateco(t),
            "capitale_sociale": extract_labeled(t, ["capitale sociale"], 70),
            "rea": extract_labeled(t, ["rea"], 50),
        })
    elif category == "Bilancio":
        meta.update({
            "anno": extract_labeled(t, ["esercizio", "anno"], 30),
            "ricavi": extract_labeled(t, ["ricavi", "valore della produzione"], 70),
            "utile_perdita": extract_labeled(t, ["utile", "perdita", "risultato esercizio"], 70),
            "patrimonio_netto": extract_labeled(t, ["patrimonio netto"], 70),
            "debiti": extract_labeled(t, ["debiti"], 70),
            "ebitda_stimato": extract_labeled(t, ["ebitda", "margine operativo lordo"], 70),
        })
    elif category == "Centrale Rischi":
        meta.update({
            "mese": extract_labeled(t, ["mese", "rilevazione"], 60),
            "banca": extract_labeled(t, ["intermediario", "banca"], 100),
            "accordato": extract_labeled(t, ["accordato"], 60),
            "utilizzato": extract_labeled(t, ["utilizzato"], 60),
            "scaduto": extract_labeled(t, ["scaduto"], 60),
            "sconfini": extract_labeled(t, ["sconfini", "sconfinamento"], 60),
            "garanzie": extract_labeled(t, ["garanzie"], 80),
        })
    elif category == "DURC":
        meta.update({
            "protocollo": extract_protocollo(t),
            "data_validita": extract_labeled(t, ["valido fino al", "data validita", "scadenza"], 40),
            "esito": extract_labeled(t, ["esito", "regolare"], 80),
        })
    elif category == "Documento di identita":
        meta.update({
            "nome": extract_labeled(t, ["nome"], 80),
            "cognome": extract_labeled(t, ["cognome"], 80),
            "scadenza": extract_labeled(t, ["scadenza", "valida fino"], 40),
            "codice_fiscale": extract_cf(t),
        })
    elif category == "Contratto":
        meta.update({
            "parti": extract_labeled(t, ["parti", "tra", "contraenti"], 160),
            "decorrenza": extract_labeled(t, ["decorrenza", "data inizio"], 60),
            "durata": extract_labeled(t, ["durata"], 70),
            "importo": extract_primary_amount(t),
            "scadenze": extract_labeled(t, ["scadenze", "scadenza"], 100),
        })
    elif category == "Fattura":
        meta.update({
            "numero_fattura": extract_labeled(t, ["numero fattura", "fattura n", "n fattura"], 50),
            "data": extract_document_date(t),
            "imponibile": extract_labeled(t, ["imponibile"], 60),
            "iva": extract_labeled(t, ["iva"], 60),
            "totale": extract_primary_amount(t),
        })
    elif category == "Estratto conto":
        meta.update({
            "banca": extract_labeled(t, ["banca", "istituto"], 100),
            "iban": extract_iban(t),
            "periodo": extract_labeled(t, ["periodo", "dal"], 100),
            "saldo_iniziale": extract_labeled(t, ["saldo iniziale"], 60),
            "saldo_finale": extract_labeled(t, ["saldo finale"], 60),
        })
    elif category == "Preventivo":
        meta.update({
            "fornitore": extract_labeled(t, ["fornitore", "mittente"], 100),
            "oggetto": extract_labeled(t, ["oggetto"], 160),
            "importo": extract_primary_amount(t),
            "validita": extract_labeled(t, ["validita", "valido fino"], 70),
        })
    elif category == "Report bancario":
        meta.update({
            "cliente": extract_labeled(t, ["cliente"], 100),
            "banca": extract_labeled(t, ["banca"], 100),
            "importo_richiesto": extract_labeled(t, ["importo richiesto", "finanziamento"], 70),
            "durata": extract_labeled(t, ["durata"], 70),
            "garanzie": extract_labeled(t, ["garanzie"], 100),
        })
    elif category == "Business plan":
        meta.update({
            "cliente": extract_labeled(t, ["cliente", "azienda"], 100),
            "investimento": extract_labeled(t, ["investimento"], 100),
            "fabbisogno": extract_labeled(t, ["fabbisogno"], 100),
            "durata": extract_labeled(t, ["durata"], 70),
            "previsioni_economiche": extract_labeled(t, ["previsioni", "forecast"], 160),
        })
    elif category == "Dichiarazione fiscale":
        meta.update({
            "contribuente": extract_labeled(t, ["contribuente"], 100),
            "anno": extract_labeled(t, ["anno", "periodo imposta"], 50),
            "imposte": extract_labeled(t, ["imposte", "saldo"], 80),
            "reddito": extract_labeled(t, ["reddito"], 80),
            "protocollo": extract_protocollo(t),
        })
    return {k: v for k, v in meta.items() if v not in (None, "", [])}


def local_classify(path: Path, text: str) -> ClassificationResult:
    scores = keyword_scores(text, path.name)
    category, score = max(scores.items(), key=lambda x: x[1]) if scores else ("Documento generico", 0.0)
    if score < 0.08:
        category = "Documento generico"
    piva = extract_piva(text)
    cf = extract_cf(text)
    ragione = extract_ragione_sociale(text, path.name)
    email = extract_email(text)
    pec = extract_pec(text)
    admin = extract_amministratore(text)
    sede = extract_sede(text)
    ateco = extract_ateco(text)
    doc_date = extract_document_date(text)
    amount = extract_primary_amount(text)
    iban = extract_iban(text)
    protocollo = extract_protocollo(text)
    learned = learned_hint(text, piva, cf, ragione)
    reasons = []
    if score > 0:
        reasons.append(f"keyword categoria {category}: {score:.2f}")
    if learned.get("row"):
        row = learned["row"]
        if row.get("categoria_corretta") in DOCUMENT_CATEGORIES:
            category = row["categoria_corretta"]
            score = max(score, 0.80)
            reasons.append("classificazione suggerita da apprendimento")
        if not ragione and row.get("cliente_corretto"):
            ragione = row["cliente_corretto"]
    client = find_client(piva, cf, ragione, email, pec, admin)
    matched_client_id = int(client["id"]) if client else None
    if client:
        ragione = ragione or client.get("ragione_sociale") or ""
        piva = piva or client.get("partita_iva") or ""
        cf = cf or client.get("codice_fiscale") or ""
        admin = admin or client.get("amministratore") or ""
        ateco = ateco or client.get("ateco") or ""
        sede = sede or client.get("sede") or ""
    client_conf = 0.0
    if piva:
        client_conf += 0.35
    if cf:
        client_conf += 0.20
    if ragione:
        client_conf += 0.22
    if matched_client_id:
        client_conf += 0.20
    if admin:
        client_conf += 0.08
    if email or pec:
        client_conf += 0.08
    category_conf = 0.20 + min(score * 1.35, 0.45)
    learn_conf = min(float(learned.get("score", 0.0)), 0.25)
    confidence = min(0.99, category_conf + client_conf + learn_conf)
    anomalies = []
    if not (piva or cf or ragione):
        anomalies.append("Cliente non riconosciuto con certezza")
    if category == "Documento generico":
        anomalies.append("Categoria non determinata")
    if not text or len(normalize_spaces(text)) < 40:
        anomalies.append("Testo estratto insufficiente")
    needs_review = confidence < 0.68 or category == "Documento generico" or not (piva or cf or ragione)
    meta = {
        "ragione_sociale": ragione,
        "partita_iva": piva,
        "codice_fiscale": cf,
        "sede": sede,
        "amministratore": admin,
        "ateco": ateco,
        "email": email,
        "pec": pec,
        "iban": iban,
        "protocollo": protocollo,
        "keywords": top_keywords(text),
        "dati_categoria": extract_category_metadata(category, text),
        "score_categoria": score,
        "score_apprendimento": learned.get("score", 0.0),
    }
    return ClassificationResult(
        categoria=category,
        confidence=round(confidence, 3),
        ragione_sociale=ragione,
        partita_iva=piva,
        codice_fiscale=cf,
        sede=sede,
        amministratore=admin,
        ateco=ateco,
        email=email,
        pec=pec,
        data_documento=doc_date,
        importo=amount,
        iban=iban,
        protocollo=protocollo,
        metadata=meta,
        motivazione="; ".join(reasons) or "Classificazione locale rule-based con pattern fiscali e parole chiave.",
        anomalie=anomalies,
        needs_review=needs_review,
        matched_client_id=matched_client_id,
    )


def ai_classify_optional(path: Path, text: str, local: ClassificationResult) -> ClassificationResult:
    if not os.getenv("OPENAI_API_KEY") or not os.getenv("OPENAI_MODEL"):
        return local
    try:
        from openai import OpenAI
        client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        prompt = {
            "file_name": path.name,
            "local_classification": asdict(local),
            "text_excerpt": (text or "")[:14000],
            "categories": DOCUMENT_CATEGORIES,
            "priority": ["ragione_sociale", "partita_iva", "codice_fiscale", "amministratore", "sede", "email", "pec"],
            "output": "JSON con categoria, confidence 0-1, dati cliente, metadati, motivazione, anomalie, needs_review",
        }
        response = client.chat.completions.create(
            model=os.getenv("OPENAI_MODEL"),
            messages=[
                {"role": "system", "content": "Sei un motore IDP per documenti fiscali, bancari e societari italiani. Rispondi solo JSON valido."},
                {"role": "user", "content": json.dumps(prompt, ensure_ascii=False)},
            ],
            temperature=0.0,
        )
        raw = response.choices[0].message.content or "{}"
        data = json.loads(raw)
        merged = asdict(local)
        mapping = {
            "categoria": "categoria",
            "confidence": "confidence",
            "ragione_sociale": "ragione_sociale",
            "partita_iva": "partita_iva",
            "codice_fiscale": "codice_fiscale",
            "sede": "sede",
            "amministratore": "amministratore",
            "ateco": "ateco",
            "email": "email",
            "pec": "pec",
            "data_documento": "data_documento",
            "importo": "importo",
            "iban": "iban",
            "protocollo": "protocollo",
            "metadata": "metadata",
            "motivazione": "motivazione",
            "anomalie": "anomalie",
            "needs_review": "needs_review",
        }
        for src, dst in mapping.items():
            if src in data and data[src] not in (None, ""):
                merged[dst] = data[src]
        if merged.get("categoria") not in DOCUMENT_CATEGORIES:
            merged["categoria"] = local.categoria
        merged["motivazione"] = "IA: " + str(merged.get("motivazione") or local.motivazione)
        result = ClassificationResult(**merged)
        result.confidence = max(float(result.confidence or 0), local.confidence)
        result.needs_review = bool(result.needs_review) or result.confidence < 0.72
        return result
    except Exception as exc:
        local.motivazione = f"{local.motivazione} | IA non usata: {exc}"
        return local


def build_target_path(src: Path, c: ClassificationResult, status: str) -> Path:
    doc_dt = parse_date_any(c.data_documento) or datetime.now()
    year, month = month_folder(doc_dt)
    ext = src.suffix.lower() or ".bin"
    if status == "da_verificare":
        base = VERIFY_DIR
        client_folder = "DA_VERIFICARE"
    else:
        ragione = c.ragione_sociale or "CLIENTE_SENZA_NOME"
        ident = c.partita_iva or c.codice_fiscale or "NO_PIVA_CF"
        client_folder = f"{safe_filename(ragione, 85)}_{safe_filename(ident, 30)}"
        base = ARCHIVE_DIR
    category_folder = CATEGORY_FOLDERS.get(c.categoria, "Altro")
    date_name = ddmmyyyy(doc_dt)
    name_client = safe_filename(c.ragione_sociale or "CLIENTE_DA_VERIFICARE", 80)
    name_cat = safe_filename(categoria_to_file_label(c.categoria), 45)
    file_name = f"{name_client}_{name_cat}_{date_name}{ext}"
    return base / client_folder / year / month / category_folder / file_name


def categoria_to_file_label(category: str) -> str:
    labels = {
        "Visura camerale": "visura",
        "Centrale Rischi": "centrale_rischi",
        "Documento di identita": "documento_identita",
        "Estratto conto": "estratto_conto",
        "Business plan": "business_plan",
        "Report bancario": "report_bancario",
        "Dichiarazione fiscale": "dichiarazione_fiscale",
        "Documento generico": "documento_generico",
    }
    return labels.get(category, category.lower().replace(" ", "_"))


def insert_document(src: Path, stored: Path, sha: str, extraction: ExtractionResult, c: ClassificationResult, status: str, client_id: Optional[int]) -> int:
    now = now_iso()
    meta = c.metadata or {}
    meta["motivazione"] = c.motivazione
    meta["anomalie"] = c.anomalie or []
    con = connect()
    try:
        cur = con.execute(
            """
            INSERT INTO documenti(cliente_id,categoria,nome_file_originale,nome_file_archiviato,percorso_file,
                data_documento,importo,hash_sha256,testo_estratto,stato_riconoscimento,confidenza,motore_estrazione,
                metadata_json,data_importazione,data_ultimo_aggiornamento)
            VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            """,
            (
                client_id, c.categoria, src.name, stored.name, str(stored), c.data_documento, c.importo, sha,
                extraction.text, status, float(c.confidence), extraction.engine, json.dumps(meta, ensure_ascii=False), now, now
            ),
        )
        doc_id = int(cur.lastrowid)
        try:
            con.execute(
                "INSERT INTO documenti_fts(documento_id,titolo,corpo,metadata) VALUES(?,?,?,?)",
                (doc_id, stored.name, extraction.text or "", json.dumps(meta, ensure_ascii=False)),
            )
        except Exception:
            pass
        con.commit()
        return doc_id
    finally:
        con.close()


def add_to_review_queue(doc_id: int, stored: Path, extraction: ExtractionResult, c: ClassificationResult) -> None:
    motivo = "; ".join(c.anomalie or []) or "Confidenza insufficiente"
    execute(
        """
        INSERT INTO coda_da_verificare(documento_id,file,testo_estratto,categoria_suggerita,cliente_suggerito,motivo_incertezza,stato_lavorazione,data_creazione)
        VALUES(?,?,?,?,?,?,?,?)
        """,
        (doc_id, str(stored), extraction.text, c.categoria, c.ragione_sociale, motivo, "aperta", now_iso()),
    )


def learn_from_confirmation(categoria: str, cliente: str, piva: str, cf: str, text: str, source: str) -> None:
    keys = ",".join(top_keywords(text))
    regola = json.dumps({"categoria": categoria, "cliente": cliente, "piva": piva, "cf": cf, "keywords": keys.split(",")}, ensure_ascii=False)
    execute(
        """
        INSERT INTO modello_apprendimento(testo_parole_chiave,categoria_corretta,cliente_corretto,partita_iva,codice_fiscale,regola_appresa,data_correzione,fonte_correzione)
        VALUES(?,?,?,?,?,?,?,?)
        """,
        (keys, categoria, cliente, piva, cf, regola, now_iso(), source),
    )


def process_file(src: Path, use_local_ocr: bool = True, use_cloud_ocr: bool = False, use_ai: bool = False, max_pages_ocr: int = 8, force_review: bool = False) -> dict[str, Any]:
    sha = sha256_file(src)
    existing = query_one("SELECT * FROM documenti WHERE hash_sha256=?", (sha,))
    if existing:
        log_event("DUPLICATO_BLOCCATO", src.name, "DUPLICATO", f"Gia presente: {existing.get('percorso_file')}")
        return {"ok": True, "duplicate": True, "message": "Duplicato bloccato", "existing": existing}
    extraction = extract_text(src, use_local_ocr=use_local_ocr, use_cloud_ocr=use_cloud_ocr, max_pages_ocr=max_pages_ocr)
    classification = local_classify(src, extraction.text)
    if use_ai:
        classification = ai_classify_optional(src, extraction.text, classification)
    if force_review:
        classification.needs_review = True
        classification.anomalie = (classification.anomalie or []) + ["Verifica forzata dall'utente"]
    status = "da_verificare" if classification.needs_review else "archiviato"
    client_id = None
    if status == "archiviato":
        client_id = ensure_client(classification, create=True)
        classification.matched_client_id = client_id
        if client_id is None:
            status = "da_verificare"
            classification.anomalie = (classification.anomalie or []) + ["Cliente non creato: dati insufficienti"]
    target = build_target_path(src, classification, status)
    target.parent.mkdir(parents=True, exist_ok=True)
    stored = unique_path(target)
    shutil.copy2(str(src), str(stored))
    doc_id = insert_document(src, stored, sha, extraction, classification, status, client_id)
    sidecar = stored.with_suffix(stored.suffix + ".idp.json")
    try:
        sidecar.write_text(json.dumps({"classification": asdict(classification), "extraction": asdict(extraction)}, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass
    if status == "da_verificare":
        add_to_review_queue(doc_id, stored, extraction, classification)
        log_event("CODA_DA_VERIFICARE", src.name, "DA_VERIFICARE", "; ".join(classification.anomalie or []))
    else:
        learn_from_confirmation(classification.categoria, classification.ragione_sociale, classification.partita_iva, classification.codice_fiscale, extraction.text, "auto_confermato")
        log_event("ARCHIVIAZIONE", src.name, "ARCHIVIATO", str(stored))
    return {
        "ok": True,
        "duplicate": False,
        "doc_id": doc_id,
        "status": status,
        "stored_path": str(stored),
        "classification": asdict(classification),
        "extraction_engine": extraction.engine,
        "warnings": extraction.warnings or [],
        "message": f"Documento {status}",
    }


def iter_supported_files(folder: Path, recursive: bool = True) -> list[Path]:
    skip = {".git", "__pycache__", ".venv", "venv", "node_modules"}
    iterator = folder.rglob("*") if recursive else folder.glob("*")
    out = []
    archive_resolved = ARCHIVE_DIR.resolve()
    for p in iterator:
        try:
            if not p.is_file() or p.suffix.lower() not in SUPPORTED_EXTENSIONS:
                continue
            if any(part in skip for part in p.parts):
                continue
            if str(p.resolve()).startswith(str(archive_resolved)):
                continue
            out.append(p)
        except Exception:
            continue
    return out


def get_dashboard_stats() -> dict[str, Any]:
    today = datetime.now().date().isoformat()
    stats = {}
    stats["total_documenti"] = query_one("SELECT COUNT(*) n FROM documenti")["n"]
    stats["clienti"] = query_one("SELECT COUNT(*) n FROM clienti")["n"]
    stats["importati_oggi"] = query_one("SELECT COUNT(*) n FROM documenti WHERE substr(data_importazione,1,10)=?", (today,))["n"]
    stats["duplicati"] = query_one("SELECT COUNT(*) n FROM log_attivita WHERE operazione='DUPLICATO_BLOCCATO'")["n"]
    stats["non_riconosciuti"] = query_one("SELECT COUNT(*) n FROM documenti WHERE stato_riconoscimento='da_verificare'")["n"]
    stats["coda"] = query_one("SELECT COUNT(*) n FROM coda_da_verificare WHERE stato_lavorazione='aperta'")["n"]
    stats["db_size_mb"] = round(DB_PATH.stat().st_size / (1024 * 1024), 2) if DB_PATH.exists() else 0
    last_backup = sorted(BACKUP_DIR.glob("*.zip"), key=lambda x: x.stat().st_mtime, reverse=True)
    stats["last_backup"] = datetime.fromtimestamp(last_backup[0].stat().st_mtime).strftime("%d-%m-%Y %H:%M") if last_backup else "Mai"
    return stats


def export_table_csv(table: str) -> bytes:
    rows = query_all(f"SELECT * FROM {table}")
    output = io.StringIO()
    if rows:
        writer = csv.DictWriter(output, fieldnames=list(rows[0].keys()), delimiter=";")
        writer.writeheader()
        writer.writerows(rows)
    return output.getvalue().encode("utf-8-sig")


def export_table_excel(table: str) -> bytes:
    rows = query_all(f"SELECT * FROM {table}")
    if pd is None:
        return export_table_csv(table)
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine="openpyxl") as writer:
        pd.DataFrame(rows).to_excel(writer, index=False, sheet_name=table[:31])
    return bio.getvalue()


def generate_pdf_report(kind: str = "generale", filter_value: str = "") -> Path:
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = REPORT_DIR / f"report_idp_{kind}_{stamp}.pdf"
    rows = query_all(
        """
        SELECT d.id, d.categoria, d.nome_file_archiviato, d.stato_riconoscimento, d.confidenza, d.data_importazione,
               c.ragione_sociale, c.partita_iva
        FROM documenti d LEFT JOIN clienti c ON c.id=d.cliente_id
        ORDER BY d.id DESC LIMIT 500
        """
    )
    stats = get_dashboard_stats()
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
        styles = getSampleStyleSheet()
        story = []
        story.append(Paragraph("FinancePlus Archive AI - Report IDP", styles["Title"]))
        story.append(Paragraph(f"Generato il {datetime.now().strftime('%d-%m-%Y %H:%M')}", styles["Normal"]))
        story.append(Spacer(1, 12))
        summary = [
            ["Indicatore", "Valore"],
            ["Documenti archiviati", str(stats["total_documenti"])],
            ["Clienti riconosciuti", str(stats["clienti"])],
            ["Importati oggi", str(stats["importati_oggi"])],
            ["Duplicati bloccati", str(stats["duplicati"])],
            ["Da verificare", str(stats["coda"])],
            ["Ultimo backup", str(stats["last_backup"])],
        ]
        t = Table(summary, colWidths=[220, 220])
        t.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0B1F3A")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("GRID", (0, 0), (-1, -1), 0.25, colors.grey), ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold")]))
        story.append(t)
        story.append(Spacer(1, 18))
        story.append(Paragraph("Ultimi documenti", styles["Heading2"]))
        data = [["ID", "Cliente", "Categoria", "Stato", "Conf.", "Data"]]
        for r in rows[:80]:
            data.append([str(r.get("id", "")), str(r.get("ragione_sociale") or "" )[:28], str(r.get("categoria") or "")[:22], str(r.get("stato_riconoscimento") or ""), str(r.get("confidenza") or ""), str(r.get("data_importazione") or "")[:10]])
        tbl = Table(data, colWidths=[35, 135, 105, 75, 45, 60])
        tbl.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#B87333")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("GRID", (0, 0), (-1, -1), 0.25, colors.lightgrey), ("FONTSIZE", (0, 0), (-1, -1), 7)]))
        story.append(tbl)
        doc.build(story)
    except Exception:
        text_path = path.with_suffix(".txt")
        text_path.write_text(json.dumps({"stats": stats, "rows": rows}, ensure_ascii=False, indent=2), encoding="utf-8")
        return text_path
    return path


def create_backup_zip(include_temp: bool = False) -> Path:
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    zip_path = BACKUP_DIR / f"backup_financeplus_archive_ai_{stamp}.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        if DB_PATH.exists():
            z.write(DB_PATH, arcname="database/financeplus_archive_ai.db")
        for root in [ARCHIVE_DIR, REPORT_DIR, LOG_DIR]:
            if not root.exists():
                continue
            for f in root.rglob("*"):
                if f.is_file():
                    z.write(f, arcname=str(f.relative_to(DATA_DIR)))
        settings = query_all("SELECT * FROM settings")
        learning = query_all("SELECT * FROM modello_apprendimento")
        z.writestr("config/settings.json", json.dumps(settings, ensure_ascii=False, indent=2))
        z.writestr("config/modello_apprendimento.json", json.dumps(learning, ensure_ascii=False, indent=2))
        if include_temp and TEMP_DIR.exists():
            for f in TEMP_DIR.rglob("*"):
                if f.is_file():
                    z.write(f, arcname=str(f.relative_to(DATA_DIR)))
    log_event("BACKUP", zip_path.name, "OK", str(zip_path))
    return zip_path


def authenticate(username: str, password: str) -> Optional[dict[str, Any]]:
    row = query_one("SELECT * FROM users WHERE username=? AND active=1", (username,))
    if row and row.get("password_hash") == hash_password(password):
        return row
    return None


def require_perm(perm: str) -> bool:
    user = st.session_state.get("user") if st else None
    if not user:
        return False
    return perm in ROLE_PERMS.get(user.get("role", "Solo lettura"), set())


def render_login() -> None:
    st.title(APP_TITLE)
    st.info("Accesso locale. Primo accesso: admin / admin123. Cambiare password dopo il primo utilizzo.")
    with st.form("login_form"):
        username = st.text_input("Utente", value="admin")
        password = st.text_input("Password", type="password")
        submitted = st.form_submit_button("Accedi", type="primary")
    if submitted:
        user = authenticate(username.strip(), password)
        if user:
            st.session_state.user = user
            log_event("LOGIN", username, "OK")
            st.rerun()
        else:
            log_event("LOGIN", username, "KO", "Credenziali errate")
            st.error("Credenziali non valide.")


def render_header() -> None:
    user = st.session_state.get("user") or {}
    st.title(APP_TITLE)
    st.caption(f"Versione {APP_VERSION} | Utente: {user.get('username')} | Ruolo: {user.get('role')}")


def render_dashboard() -> None:
    render_header()
    stats = get_dashboard_stats()
    c1, c2, c3, c4, c5, c6 = st.columns(6)
    c1.metric("Documenti", stats["total_documenti"])
    c2.metric("Clienti", stats["clienti"])
    c3.metric("Importati oggi", stats["importati_oggi"])
    c4.metric("Duplicati bloccati", stats["duplicati"])
    c5.metric("Non riconosciuti", stats["non_riconosciuti"])
    c6.metric("Da verificare", stats["coda"])
    st.divider()
    s1, s2, s3 = st.columns(3)
    s1.success(f"Database SQLite: OK - {stats['db_size_mb']} MB")
    s2.info(f"Backup ultimo: {stats['last_backup']}")
    s3.write(f"Archivio: `{ARCHIVE_DIR}`")
    col1, col2 = st.columns([1, 1])
    with col1:
        st.subheader("Statistiche per categoria")
        rows = query_all("SELECT categoria, COUNT(*) totale FROM documenti GROUP BY categoria ORDER BY totale DESC")
        if rows and pd:
            st.dataframe(pd.DataFrame(rows), width="stretch", hide_index=True)
            st.bar_chart(pd.DataFrame(rows).set_index("categoria"))
        elif rows:
            st.table(rows)
        else:
            st.info("Nessun documento presente.")
    with col2:
        st.subheader("Ultime attivita")
        logs = query_all("SELECT data, utente, operazione, documento, risultato FROM log_attivita ORDER BY id DESC LIMIT 15")
        if logs and pd:
            st.dataframe(pd.DataFrame(logs), width="stretch", hide_index=True)
        elif logs:
            st.table(logs)
        else:
            st.info("Log vuoto.")


def save_uploaded_file(uploaded) -> Path:
    TEMP_DIR.mkdir(parents=True, exist_ok=True)
    safe = safe_filename(uploaded.name, 120)
    path = unique_path(TEMP_DIR / safe)
    with path.open("wb") as f:
        f.write(uploaded.getbuffer())
    return path


def choose_folder_dialog(initial_dir: str | None = None) -> str:
    """Apre il selettore cartella nativo quando Streamlit gira sul PC locale.

    Nota: il dialog funziona su Windows/Mac/Linux con ambiente grafico. Se Streamlit
    gira su server remoto o in ambiente senza desktop, usare il campo percorso.
    """
    try:
        import tkinter as tk
        from tkinter import filedialog
        root = tk.Tk()
        root.withdraw()
        try:
            root.wm_attributes("-topmost", 1)
        except Exception:
            pass
        selected = filedialog.askdirectory(
            initialdir=initial_dir or str(Path.home()),
            title="Seleziona cartella documenti FinancePlus",
            mustexist=True,
        )
        root.destroy()
        return selected or ""
    except Exception as exc:
        if st:
            st.warning(f"Selettore cartella non disponibile in questo ambiente. Usa il percorso manuale. Dettaglio: {exc}")
        return ""


def common_local_folders() -> dict[str, str]:
    home = Path.home()
    folders: dict[str, str] = {}
    candidates = [
        ("Desktop", home / "Desktop"),
        ("Documenti", home / "Documents"),
        ("Download", home / "Downloads"),
    ]
    onedrive = os.environ.get("OneDrive") or os.environ.get("OneDriveCommercial") or os.environ.get("OneDriveConsumer")
    if onedrive:
        candidates.append(("OneDrive", Path(onedrive)))
    for label, path in candidates:
        try:
            if path.exists() and path.is_dir():
                folders[label] = str(path)
        except Exception:
            pass
    return folders


def render_import_summary(results: list[dict[str, Any]]) -> None:
    if not results:
        return
    st.divider()
    st.subheader("Risultati ultima importazione")
    imported = sum(1 for r in results if r.get("ok") and not r.get("duplicate"))
    duplicated = sum(1 for r in results if r.get("duplicate"))
    review = sum(1 for r in results if r.get("status") == "da_verificare")
    errors = sum(1 for r in results if not r.get("ok"))
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Importati/Copiati", imported)
    c2.metric("Duplicati bloccati", duplicated)
    c3.metric("Da verificare", review)
    c4.metric("Errori", errors)
    flat = []
    for r in results:
        c = r.get("classification") or {}
        flat.append({
            "ok": r.get("ok"),
            "duplicato": r.get("duplicate", False),
            "stato": r.get("status") or ("duplicato" if r.get("duplicate") else "errore"),
            "doc_id": r.get("doc_id"),
            "categoria": c.get("categoria"),
            "cliente": c.get("ragione_sociale"),
            "piva": c.get("partita_iva"),
            "confidenza": c.get("confidence"),
            "percorso": r.get("stored_path") or (r.get("existing") or {}).get("percorso_file"),
            "messaggio": r.get("message") or r.get("error"),
        })
    if pd:
        st.dataframe(pd.DataFrame(flat), width="stretch", hide_index=True)
    else:
        st.table(flat)


def render_import() -> None:
    render_header()
    if not require_perm("import"):
        st.error("Ruolo non autorizzato all'importazione.")
        return

    st.subheader("Importa e cataloga documenti")
    st.info(
        "Modalita consigliata: seleziona una CARTELLA. Il sistema legge automaticamente anche tutte le sottocartelle, "
        "trova PDF/immagini/TXT/documenti supportati, controlla i duplicati e archivia tutto per cliente, mese e categoria."
    )

    with st.expander("Funzioni attive", expanded=False):
        st.markdown(
            """
            - Scansione cartella locale con sottocartelle.
            - Lettura PDF digitali con PyMuPDF.
            - OCR locale con pytesseract su immagini e PDF scannerizzati.
            - Predisposizione OCR cloud/HTR senza invio automatico dei dati.
            - Classificazione IDP rule-based e opzionale IA.
            - Riconoscimento cliente da ragione sociale, P.IVA, CF, amministratore, email/PEC e anagrafica.
            - Hash SHA256 anti-duplicati, rinomina e archivio per cliente / anno / mese / categoria.
            """
        )

    opt1, opt2, opt3, opt4 = st.columns(4)
    use_local_ocr = opt1.checkbox("OCR locale", value=True)
    use_cloud_ocr = opt2.checkbox("Predisposizione OCR cloud/HTR", value=False)
    use_ai = opt3.checkbox("Usa IA se configurata", value=False)
    force_review = opt4.checkbox("Forza coda verifica", value=False)
    max_pages_ocr = st.slider("Massimo pagine OCR per PDF scannerizzato", 1, 50, 8)

    tab_folder, tab_zip, tab_files = st.tabs([
        "1. Seleziona cartella e sottocartelle",
        "2. Importa cartella ZIP",
        "3. File singoli",
    ])

    with tab_folder:
        st.markdown("### Scansione massiva da cartella")
        st.caption("Funziona quando Streamlit gira sul tuo PC. Puoi usare il pulsante Sfoglia oppure incollare il percorso della cartella.")

        if "folder_path" not in st.session_state:
            st.session_state["folder_path"] = ""

        quick = common_local_folders()
        q1, q2 = st.columns([2, 1])
        with q1:
            if quick:
                selected_quick = st.selectbox("Scelta rapida", ["- Nessuna -"] + list(quick.keys()))
                if selected_quick != "- Nessuna -":
                    st.session_state["folder_path"] = quick[selected_quick]
        with q2:
            if st.button("Sfoglia cartella Windows", type="secondary"):
                picked = choose_folder_dialog(st.session_state.get("folder_path") or None)
                if picked:
                    st.session_state["folder_path"] = picked
                    st.rerun()

        folder_raw = st.text_input(
            "Cartella da importare",
            key="folder_path",
            placeholder=r"Esempio: C:\Users\dani\Desktop\Documenti_Clienti oppure C:\Users\dani\OneDrive\Desktop\FP_IDP_PRO",
        )
        recursive = st.checkbox("Includi tutte le sottocartelle", value=True)
        only_new = st.checkbox("Salta automaticamente i duplicati gia archiviati", value=True, disabled=True)

        folder_clean = (folder_raw or "").strip().strip('"')
        if folder_clean:
            folder_clean = os.path.expandvars(folder_clean)
            p = Path(folder_clean).expanduser()
            if p.exists() and p.is_dir():
                with st.spinner("Scansione cartella in corso..."):
                    found = iter_supported_files(p, recursive)
                st.success(f"Cartella valida. File documentali trovati: {len(found)}")

                col_a, col_b, col_c = st.columns(3)
                col_a.write(f"**Percorso:** `{p}`")
                col_b.write(f"**Sottocartelle:** {'Si' if recursive else 'No'}")
                col_c.write(f"**Estensioni:** {', '.join(sorted(SUPPORTED_EXTENSIONS))}")

                with st.expander("Anteprima dei file trovati", expanded=False):
                    preview = [{"file": f.name, "cartella": str(f.parent), "estensione": f.suffix.lower()} for f in found[:500]]
                    if preview and pd:
                        st.dataframe(pd.DataFrame(preview), width="stretch", hide_index=True)
                    elif preview:
                        st.table(preview)
                    else:
                        st.info("Nessun file supportato trovato in questa cartella.")
                    if len(found) > 500:
                        st.caption(f"Mostrati i primi 500 file su {len(found)}.")

                if st.button("Scansiona cartella e cataloga TUTTE le sottocartelle", type="primary", disabled=not found):
                    results: list[dict[str, Any]] = []
                    progress = st.progress(0)
                    status_box = st.empty()
                    for idx, f in enumerate(found):
                        try:
                            status_box.write(f"Elaborazione {idx + 1}/{len(found)}: `{f.name}`")
                            results.append(process_file(f, use_local_ocr, use_cloud_ocr, use_ai, max_pages_ocr, force_review))
                        except Exception as exc:
                            results.append({"ok": False, "file": str(f), "error": str(exc), "trace": traceback.format_exc()})
                            log_event("ERRORE_IMPORT_CARTELLA", str(f), "KO", str(exc))
                        progress.progress((idx + 1) / max(1, len(found)))
                    status_box.empty()
                    st.session_state["last_import_results"] = results
                    st.success(f"Scansione conclusa: {len(results)} file elaborati dalla cartella e dalle sottocartelle.")
            else:
                st.warning("Percorso non valido o cartella non accessibile. Controlla che il percorso esista sul PC dove hai avviato Streamlit.")
        else:
            st.info("Seleziona o incolla una cartella per avviare la scansione massiva.")

    with tab_zip:
        st.markdown("### Importa una cartella compressa ZIP")
        st.caption("Utile se lavori da browser o vuoi mantenere la struttura con sottocartelle. Comprimi la cartella in .zip e caricala qui.")
        zip_file = st.file_uploader("Carica ZIP della cartella documenti", type=["zip"], key="zip_folder_upload")
        if st.button("Estrai ZIP e cataloga tutti i documenti", type="primary", disabled=not zip_file):
            results: list[dict[str, Any]] = []
            extract_dir = TEMP_DIR / f"zip_import_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            extract_dir.mkdir(parents=True, exist_ok=True)
            try:
                zip_path = save_uploaded_file(zip_file)
                with zipfile.ZipFile(zip_path, "r") as zf:
                    zf.extractall(extract_dir)
                found = iter_supported_files(extract_dir, recursive=True)
                progress = st.progress(0)
                status_box = st.empty()
                for idx, f in enumerate(found):
                    try:
                        status_box.write(f"Elaborazione {idx + 1}/{len(found)}: `{f.name}`")
                        results.append(process_file(f, use_local_ocr, use_cloud_ocr, use_ai, max_pages_ocr, force_review))
                    except Exception as exc:
                        results.append({"ok": False, "file": str(f), "error": str(exc), "trace": traceback.format_exc()})
                        log_event("ERRORE_IMPORT_ZIP", str(f), "KO", str(exc))
                    progress.progress((idx + 1) / max(1, len(found)))
                status_box.empty()
                st.session_state["last_import_results"] = results
                st.success(f"ZIP elaborato: {len(results)} documenti processati.")
            except Exception as exc:
                st.error(f"Errore import ZIP: {exc}")
                log_event("ERRORE_IMPORT_ZIP", zip_file.name if zip_file else "ZIP", "KO", str(exc))

    with tab_files:
        st.markdown("### Caricamento manuale file singoli")
        st.caption("Usalo solo per prove o casi isolati. Per il lavoro normale usa la cartella.")
        files = st.file_uploader(
            "Carica PDF, immagini, TXT, DOCX, XLSX",
            type=[x.replace(".", "") for x in SUPPORTED_EXTENSIONS],
            accept_multiple_files=True,
            key="single_files_upload",
        )
        if st.button("Elabora file caricati", type="primary", disabled=not files):
            results: list[dict[str, Any]] = []
            progress = st.progress(0)
            for idx, uploaded in enumerate(files):
                try:
                    temp_path = save_uploaded_file(uploaded)
                    res = process_file(temp_path, use_local_ocr, use_cloud_ocr, use_ai, max_pages_ocr, force_review)
                    results.append(res)
                except Exception as exc:
                    results.append({"ok": False, "file": uploaded.name, "error": str(exc), "trace": traceback.format_exc()})
                    log_event("ERRORE_IMPORT", uploaded.name, "KO", str(exc))
                progress.progress((idx + 1) / max(1, len(files)))
            st.session_state["last_import_results"] = results
            st.success(f"Elaborazione conclusa: {len(results)} file.")

    render_import_summary(st.session_state.get("last_import_results") or [])

# -----------------------------------------------------------------------------
# MODULO EMAIL AZIENDA - IMAP, CERCA AZIENDA, VEDI TUTTO, SCARICA TUTTO
# -----------------------------------------------------------------------------

ITALIAN_MONTHS_UPPER = {
    1: "GENNAIO", 2: "FEBBRAIO", 3: "MARZO", 4: "APRILE", 5: "MAGGIO", 6: "GIUGNO",
    7: "LUGLIO", 8: "AGOSTO", 9: "SETTEMBRE", 10: "OTTOBRE", 11: "NOVEMBRE", 12: "DICEMBRE"
}


def decode_mime_header(value: str | None) -> str:
    if not value:
        return ""
    try:
        return str(make_header(decode_header(value)))
    except Exception:
        return str(value)


def parse_email_date(value: str | None) -> datetime:
    if not value:
        return datetime.now()
    try:
        dt = parsedate_to_datetime(value)
        if dt.tzinfo is not None:
            dt = dt.astimezone().replace(tzinfo=None)
        return dt
    except Exception:
        return datetime.now()


def email_date_suffix(dt: datetime | None) -> str:
    dt = dt or datetime.now()
    return f"{dt.day:02d}_{ITALIAN_MONTHS_UPPER.get(dt.month, str(dt.month))}_{dt.year}"


def normalize_search_text(value: str) -> str:
    value = value or ""
    value = value.lower()
    value = re.sub(r"[^a-z0-9àèéìòùç@._ -]+", " ", value)
    value = re.sub(r"\s+", " ", value).strip()
    return value


def email_address_folder(value: str) -> str:
    value = decode_mime_header(value or "")
    m = EMAIL_RE.search(value)
    return safe_filename((m.group(0) if m else value or "mittente_sconosciuto"), 90).lower()


def company_folder_name(value: str) -> str:
    cleaned = safe_filename(value or "AZIENDA", 100)
    return cleaned.upper()


def strip_html_basic(html_text: str) -> str:
    html_text = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", html_text or "")
    html_text = re.sub(r"(?is)<br\s*/?>", "\n", html_text)
    html_text = re.sub(r"(?is)</p>", "\n", html_text)
    html_text = re.sub(r"(?is)<.*?>", " ", html_text)
    import html as _html
    return clean_text(_html.unescape(html_text), 80000)


def extract_email_body_and_attachments(msg: EmailMessage) -> tuple[str, list[dict[str, Any]]]:
    body_parts: list[str] = []
    html_parts: list[str] = []
    attachments: list[dict[str, Any]] = []
    idx = 0
    for part in msg.walk():
        if part.is_multipart():
            continue
        ctype = (part.get_content_type() or "").lower()
        disp = (part.get_content_disposition() or "").lower()
        filename = decode_mime_header(part.get_filename() or "")
        try:
            payload_bytes = part.get_payload(decode=True) or b""
        except Exception:
            payload_bytes = b""
        if filename or disp == "attachment":
            idx += 1
            if not filename:
                ext = mimetype_to_extension(ctype)
                filename = f"allegato_{idx}{ext}"
            attachments.append({
                "filename": filename,
                "content_type": ctype,
                "size": len(payload_bytes),
                "content_disposition": disp,
                "content_id": part.get("Content-ID", ""),
                "bytes": payload_bytes,
            })
            continue
        if ctype == "text/plain":
            try:
                body_parts.append(part.get_content())
            except Exception:
                try:
                    body_parts.append(payload_bytes.decode(part.get_content_charset() or "utf-8", errors="ignore"))
                except Exception:
                    pass
        elif ctype == "text/html":
            try:
                html_parts.append(part.get_content())
            except Exception:
                try:
                    html_parts.append(payload_bytes.decode(part.get_content_charset() or "utf-8", errors="ignore"))
                except Exception:
                    pass
    body = "\n".join(x for x in body_parts if x).strip()
    if not body and html_parts:
        body = "\n".join(strip_html_basic(x) for x in html_parts if x).strip()
    return clean_text(body, 100000), attachments


def mimetype_to_extension(ctype: str) -> str:
    mapping = {
        "application/pdf": ".pdf",
        "image/jpeg": ".jpg",
        "image/png": ".png",
        "image/tiff": ".tif",
        "text/plain": ".txt",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
    }
    return mapping.get((ctype or "").lower(), ".bin")


def is_inline_or_logo_attachment(att: dict[str, Any]) -> bool:
    name = (att.get("filename") or "").lower()
    ctype = (att.get("content_type") or "").lower()
    disp = (att.get("content_disposition") or "").lower()
    size = int(att.get("size") or 0)
    if disp == "inline":
        return True
    logo_words = ["logo", "image001", "image002", "firma", "signature", "facebook", "linkedin", "instagram", "banner"]
    if ctype.startswith("image/") and (size < 120000 or any(w in name for w in logo_words)):
        return True
    return False


def summarize_email_local(subject: str, body: str, attachments: list[dict[str, Any]], company: str) -> str:
    body_clean = clean_text(re.sub(r"\s+", " ", body or " "), 4000)
    company_norm = normalize_search_text(company)
    sentences = re.split(r"(?<=[.!?])\s+", body_clean)
    picked = []
    for sentence in sentences:
        if company_norm and company_norm in normalize_search_text(sentence):
            picked.append(sentence.strip())
        if len(picked) >= 2:
            break
    if not picked:
        picked = [s.strip() for s in sentences[:2] if s.strip()]
    att_names = ", ".join((a.get("filename") or "") for a in attachments[:5])
    base = " ".join(picked)[:500] if picked else "Sintesi locale: contenuto testuale non disponibile o molto breve."
    if att_names:
        base += f" Allegati rilevati: {att_names}."
    if subject:
        base = f"Oggetto: {subject}. " + base
    return clean_text(base, 900)


def score_company_match(company: str, subject: str, body: str, attachments: list[dict[str, Any]], sender: str) -> tuple[float, str]:
    q = normalize_search_text(company)
    if not q:
        return 0.0, "azienda non indicata"
    q_tokens = [t for t in re.split(r"\s+", q) if len(t) >= 3]
    subject_n = normalize_search_text(subject)
    body_n = normalize_search_text(body[:20000])
    sender_n = normalize_search_text(sender)
    att_n = normalize_search_text(" ".join(a.get("filename", "") for a in attachments))
    score = 0.0
    reasons = []
    if q in subject_n:
        score += 0.55
        reasons.append("nome azienda nell'oggetto")
    if q in body_n:
        score += 0.30
        reasons.append("nome azienda nel corpo mail")
    if q in att_n:
        score += 0.35
        reasons.append("nome azienda negli allegati")
    if q in sender_n:
        score += 0.15
        reasons.append("nome azienda nel mittente")
    if q_tokens:
        subject_hits = sum(1 for t in q_tokens if t in subject_n)
        body_hits = sum(1 for t in q_tokens if t in body_n)
        att_hits = sum(1 for t in q_tokens if t in att_n)
        token_score = min(0.25, (subject_hits * 0.08) + (body_hits * 0.03) + (att_hits * 0.06))
        if token_score:
            score += token_score
            reasons.append("parole chiave azienda trovate")
    ratio = max(
        SequenceMatcher(None, q, subject_n[:200]).ratio() if subject_n else 0,
        SequenceMatcher(None, q, att_n[:300]).ratio() if att_n else 0,
    )
    if ratio > 0.60:
        score += min(0.20, ratio * 0.20)
        reasons.append("somiglianza testuale elevata")
    return min(score, 1.0), "; ".join(dict.fromkeys(reasons)) or "abbinamento debole"


def imap_connect(host: str, user: str, password: str, port: int = 993, use_ssl: bool = True) -> imaplib.IMAP4:
    if use_ssl:
        imap = imaplib.IMAP4_SSL(host, int(port or 993))
    else:
        imap = imaplib.IMAP4(host, int(port or 143))
    imap.login(user, password)
    return imap


def parse_mailboxes(raw: str) -> list[str]:
    boxes = [x.strip() for x in re.split(r"[,;\n]+", raw or "") if x.strip()]
    return boxes or ["INBOX"]


def fetch_company_emails(
    host: str,
    user: str,
    password: str,
    company: str,
    mailboxes: list[str],
    threshold: float = 0.25,
    max_per_mailbox: int = 150,
    port: int = 993,
    use_ssl: bool = True,
) -> list[dict[str, Any]]:
    results: list[dict[str, Any]] = []
    imap = imap_connect(host, user, password, port=port, use_ssl=use_ssl)
    try:
        for mailbox in mailboxes:
            try:
                status, _ = imap.select(f'"{mailbox}"', readonly=True)
                if status != "OK":
                    status, _ = imap.select(mailbox, readonly=True)
                if status != "OK":
                    log_event("EMAIL_MAILBOX_NON_ACCESSIBILE", mailbox, "KO")
                    continue
                status, data = imap.search(None, "ALL")
                if status != "OK" or not data or not data[0]:
                    continue
                ids = data[0].split()
                selected_ids = ids[-int(max_per_mailbox or 150):]
                for msg_id in reversed(selected_ids):
                    status, msg_data = imap.fetch(msg_id, "(RFC822)")
                    if status != "OK" or not msg_data:
                        continue
                    raw = b""
                    for item in msg_data:
                        if isinstance(item, tuple) and item[1]:
                            raw = item[1]
                            break
                    if not raw:
                        continue
                    msg = email.message_from_bytes(raw, policy=policy.default)
                    subject = decode_mime_header(msg.get("Subject", ""))
                    sender = decode_mime_header(msg.get("From", ""))
                    recipients = decode_mime_header(msg.get("To", ""))
                    dt = parse_email_date(msg.get("Date", ""))
                    message_id = str(msg.get("Message-ID", ""))
                    body, attachments = extract_email_body_and_attachments(msg)
                    score, reason = score_company_match(company, subject, body, attachments, sender)
                    if score >= float(threshold or 0.25):
                        results.append({
                            "mailbox": mailbox,
                            "imap_id": msg_id.decode("ascii", errors="ignore"),
                            "subject": subject,
                            "from": sender,
                            "to": recipients,
                            "date": dt.isoformat(),
                            "message_id": message_id,
                            "body": body,
                            "attachments": attachments,
                            "score": score,
                            "reason": reason,
                            "summary": summarize_email_local(subject, body, attachments, company),
                            "raw_bytes": raw,
                        })
            except Exception as exc:
                log_event("EMAIL_ERRORE_MAILBOX", mailbox, "KO", str(exc))
                continue
    finally:
        try:
            imap.logout()
        except Exception:
            pass
    results.sort(key=lambda r: r.get("date", ""), reverse=True)
    return results


def email_pdf_bytes(company: str, mail: dict[str, Any], saved_attachments: list[str]) -> bytes:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
        styles = getSampleStyleSheet()
        story = []
        story.append(Paragraph("FINANCEPLUS - STAMPA PDF EMAIL", styles["Title"]))
        story.append(Spacer(1, 12))
        rows = [
            ["Azienda", company],
            ["Mittente", mail.get("from", "")],
            ["Destinatari", mail.get("to", "")],
            ["Data ricezione", ddmmyyyy(parse_date_any(mail.get("date")[:10]) or datetime.fromisoformat(mail.get("date"))) if mail.get("date") else ""],
            ["Oggetto", mail.get("subject", "")],
            ["Punteggio", f"{float(mail.get('score') or 0):.0%}"],
        ]
        table = Table(rows, colWidths=[100, 360])
        table.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#eef3f8")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(table)
        story.append(Spacer(1, 14))
        story.append(Paragraph("Sintesi intelligente", styles["Heading2"]))
        story.append(Paragraph(clean_text(mail.get("summary", ""), 2000).replace("\n", "<br/>"), styles["BodyText"]))
        story.append(Spacer(1, 10))
        story.append(Paragraph("Lista allegati salvati", styles["Heading2"]))
        if saved_attachments:
            for a in saved_attachments:
                story.append(Paragraph(f"- {Path(a).name}", styles["BodyText"]))
        else:
            story.append(Paragraph("Nessun allegato salvato.", styles["BodyText"]))
        story.append(Spacer(1, 10))
        story.append(Paragraph("Contenuto email", styles["Heading2"]))
        body = clean_text(mail.get("body", ""), 9000).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
        story.append(Paragraph(body or "Contenuto testuale non disponibile.", styles["BodyText"]))
        doc.build(story)
        return buffer.getvalue()
    except Exception:
        lines = [
            "FINANCEPLUS - STAMPA PDF EMAIL",
            f"Azienda: {company}",
            f"Mittente: {mail.get('from','')}",
            f"Data: {mail.get('date','')}",
            f"Oggetto: {mail.get('subject','')}",
            "",
            "Sintesi:", mail.get("summary", ""),
            "",
            "Allegati:", "\n".join(Path(a).name for a in saved_attachments),
            "",
            "Contenuto:", mail.get("body", ""),
        ]
        return "\n".join(lines).encode("utf-8")


def register_email_file_hash(file_path: Path, company: str, sender_folder: str, source: str) -> tuple[bool, str]:
    sha = sha256_file(file_path)
    existing = query_one("SELECT * FROM email_file_hash WHERE hash_sha256=?", (sha,))
    if existing:
        try:
            file_path.unlink(missing_ok=True)
        except Exception:
            pass
        return False, str(existing.get("percorso_file"))
    execute(
        "INSERT INTO email_file_hash(hash_sha256,percorso_file,azienda,mittente,nome_file,fonte,data_archiviazione) VALUES(?,?,?,?,?,?,?)",
        (sha, str(file_path), company, sender_folder, file_path.name, source, now_iso()),
    )
    return True, str(file_path)


def save_company_email(mail: dict[str, Any], company: str, archive_root: str | None = None, skip_inline_logo: bool = True) -> dict[str, Any]:
    company_display = company.strip() or "AZIENDA"
    ensure_client(ClassificationResult(categoria="Documento generico", confidence=0.80, ragione_sociale=company_display, needs_review=False), create=True)
    raw = mail.get("raw_bytes") or b""
    message_hash = hashlib.sha256(raw).hexdigest() if raw else hashlib.sha256(json.dumps(mail, default=str).encode("utf-8")).hexdigest()
    existing = query_one("SELECT * FROM email_archiviate WHERE hash_message=?", (message_hash,))
    if existing:
        log_event("EMAIL_DUPLICATA_BLOCCATA", mail.get("subject", ""), "DUPLICATO", existing.get("pdf_path", ""))
        return {"ok": True, "duplicate": True, "message": "Email duplicata gia archiviata", "existing": existing}

    sender_folder = email_address_folder(mail.get("from", ""))
    company_folder = company_folder_name(company_display)
    root = Path(archive_root).expanduser() if archive_root else EMAIL_ARCHIVE_DIR
    target_dir = root / sender_folder / company_folder
    target_dir.mkdir(parents=True, exist_ok=True)
    dt = datetime.fromisoformat(mail.get("date")) if mail.get("date") else datetime.now()
    suffix = email_date_suffix(dt)
    saved_attachments: list[str] = []
    duplicated: list[str] = []

    for att in mail.get("attachments", []) or []:
        if skip_inline_logo and is_inline_or_logo_attachment(att):
            continue
        raw_name = att.get("filename") or "allegato.bin"
        src_name = safe_filename(Path(raw_name).stem, 100) or "allegato"
        ext = Path(raw_name).suffix or mimetype_to_extension(att.get("content_type", ""))
        target = unique_path(target_dir / f"{src_name}_{suffix}{ext}")
        try:
            target.write_bytes(att.get("bytes") or b"")
            saved, where = register_email_file_hash(target, company_folder, sender_folder, "allegato_email")
            if saved:
                saved_attachments.append(str(target))
            else:
                duplicated.append(where)
        except Exception as exc:
            log_event("EMAIL_ERRORE_SALVA_ALLEGATO", raw_name, "KO", str(exc))

    pdf_name = unique_path(target_dir / f"MAIL_{suffix}.pdf")
    pdf_content = email_pdf_bytes(company_display, mail, saved_attachments)
    pdf_name.write_bytes(pdf_content)
    try:
        register_email_file_hash(pdf_name, company_folder, sender_folder, "pdf_email")
    except Exception:
        pass

    execute(
        """
        INSERT INTO email_archiviate(azienda,mittente,destinatari,data_email,oggetto,message_id,mailbox,pdf_path,allegati_json,hash_message,score_abbinamento,motivo_abbinamento,data_archiviazione)
        VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)
        """,
        (
            company_folder, mail.get("from", ""), mail.get("to", ""), mail.get("date", ""), mail.get("subject", ""),
            mail.get("message_id", ""), mail.get("mailbox", ""), str(pdf_name), json.dumps(saved_attachments, ensure_ascii=False),
            message_hash, float(mail.get("score") or 0), mail.get("reason", ""), now_iso()
        ),
    )
    log_event("EMAIL_ARCHIVIATA", mail.get("subject", ""), "OK", str(target_dir))
    return {
        "ok": True,
        "duplicate": False,
        "pdf": str(pdf_name),
        "folder": str(target_dir),
        "attachments": saved_attachments,
        "duplicates": duplicated,
        "message": f"Email archiviata in {target_dir}",
    }


def render_email_results(results: list[dict[str, Any]], company: str) -> None:
    st.subheader("VEDI TUTTO - anteprima e sintesi intelligente")
    if not results:
        st.info("Nessuna email trovata con la soglia indicata.")
        return
    st.caption(f"Email trovate: {len(results)}. Puoi selezionare una o piu email e scaricarle.")
    for idx, mail in enumerate(results):
        dt = datetime.fromisoformat(mail.get("date")) if mail.get("date") else datetime.now()
        title = f"{idx + 1}. {ddmmyyyy(dt)} - {mail.get('subject','(senza oggetto)')} - score {float(mail.get('score') or 0):.0%}"
        with st.container(border=True):
            st.markdown(f"**{title}**")
            st.write(f"Mittente: `{mail.get('from','')}`  |  Allegati: **{len(mail.get('attachments') or [])}**")
            st.write(f"Motivo abbinamento: {mail.get('reason','')}")
            st.write(f"Sintesi: {mail.get('summary','')}")
            st.checkbox("Seleziona", key=f"email_select_{idx}")
            with st.expander("Anteprima contenuto e allegati", expanded=False):
                st.text_area("Corpo email", value=clean_text(mail.get("body", ""), 6000), height=220, key=f"email_body_{idx}")
                atts = [{"file": a.get("filename"), "tipo": a.get("content_type"), "dimensione_kb": round((a.get("size") or 0)/1024, 1)} for a in (mail.get("attachments") or [])]
                if atts and pd:
                    st.dataframe(pd.DataFrame(atts), width="stretch", hide_index=True)
                elif atts:
                    st.table(atts)
                else:
                    st.info("Nessun allegato presente.")


def render_email_company() -> None:
    render_header()
    if not require_perm("import"):
        st.error("Ruolo non autorizzato all'archiviazione email.")
        return

    st.subheader("CERCA AZIENDA - Email, allegati, anteprima e archiviazione")
    st.info(
        "Scrivi il nome azienda: il sistema cerca in oggetto, corpo della mail e nomi allegati. "
        "Con VEDI TUTTO controlli anteprima e sintesi; con SCARICA TUTTO archivi mail PDF e allegati nella cartella mittente/azienda."
    )

    with st.expander("Configurazione IMAP", expanded=True):
        c1, c2, c3 = st.columns([1.4, 0.6, 0.6])
        host = c1.text_input("Server IMAP", value=get_setting("email_imap_host", "imap.gmail.com"))
        port = c2.number_input("Porta", min_value=1, max_value=65535, value=int(get_setting("email_imap_port", "993") or 993))
        use_ssl = c3.checkbox("SSL", value=get_setting("email_imap_ssl", "1") != "0")
        c4, c5 = st.columns(2)
        email_user = c4.text_input("Email", value=get_setting("email_imap_user", ""))
        password = c5.text_input("Password / app password", value=os.environ.get("FP_EMAIL_PASSWORD", ""), type="password")
        mailboxes_raw = st.text_input("Mailbox da cercare", value=get_setting("email_mailboxes", "INBOX"), help="Esempio: INBOX oppure INBOX,Archivio")
        archive_root = st.text_input("Cartella archivio email aziende", value=get_setting("email_archive_root", str(EMAIL_ARCHIVE_DIR)))
        c6, c7, c8 = st.columns(3)
        threshold = c6.slider("Soglia riconoscimento azienda", 0.05, 1.0, float(get_setting("email_threshold", "0.25") or 0.25), 0.05)
        max_per_box = c7.number_input("Email recenti da analizzare per mailbox", min_value=10, max_value=5000, value=int(get_setting("email_max_per_box", "200") or 200), step=50)
        skip_inline = c8.checkbox("Non salvare immagini inline/logo", value=get_setting("email_skip_inline", "1") != "0")
        if st.button("Salva impostazioni email senza password"):
            set_setting("email_imap_host", host)
            set_setting("email_imap_port", str(int(port)))
            set_setting("email_imap_ssl", "1" if use_ssl else "0")
            set_setting("email_imap_user", email_user)
            set_setting("email_mailboxes", mailboxes_raw)
            set_setting("email_archive_root", archive_root)
            set_setting("email_threshold", str(threshold))
            set_setting("email_max_per_box", str(int(max_per_box)))
            set_setting("email_skip_inline", "1" if skip_inline else "0")
            st.success("Impostazioni salvate. La password non viene salvata nel database.")

    st.markdown("### Comando CERCA AZIENDA")
    company = st.text_input("Scrivi il nome azienda", placeholder="Esempio: BelGarden, PELCOM SRL, BEL GARDEN EUROPE SRL")
    b1, b2, b3 = st.columns([1, 1, 2])
    do_view = b1.button("VEDI TUTTO", type="primary", disabled=not company or not host or not email_user or not password)
    do_download_all = b2.button("SCARICA TUTTO", disabled=not company or not host or not email_user or not password)
    if not password:
        b3.caption("Inserisci una app password IMAP. Per Gmail serve una app password con IMAP attivo.")

    if do_view or do_download_all:
        with st.spinner("Ricerca email azienda in corso..."):
            try:
                results = fetch_company_emails(
                    host=host,
                    user=email_user,
                    password=password,
                    company=company,
                    mailboxes=parse_mailboxes(mailboxes_raw),
                    threshold=threshold,
                    max_per_mailbox=int(max_per_box),
                    port=int(port),
                    use_ssl=use_ssl,
                )
                st.session_state["email_company_results"] = results
                st.session_state["email_company_query"] = company
                log_event("EMAIL_CERCA_AZIENDA", company, "OK", f"Risultati: {len(results)}")
            except Exception as exc:
                st.error(f"Errore ricerca IMAP: {exc}")
                log_event("EMAIL_CERCA_AZIENDA", company, "KO", str(exc))
                results = []
        if do_download_all and results:
            saved = []
            progress = st.progress(0)
            for i, mail in enumerate(results):
                saved.append(save_company_email(mail, company, archive_root=archive_root, skip_inline_logo=skip_inline))
                progress.progress((i + 1) / max(1, len(results)))
            st.session_state["email_last_save_results"] = saved
            st.success(f"SCARICA TUTTO completato: {len(saved)} email elaborate.")

    results = st.session_state.get("email_company_results") or []
    render_email_results(results, st.session_state.get("email_company_query") or company)

    if results:
        if st.button("SCARICA EMAIL SELEZIONATE", type="primary"):
            selected = [mail for idx, mail in enumerate(results) if st.session_state.get(f"email_select_{idx}")]
            if not selected:
                st.warning("Nessuna email selezionata.")
            else:
                saved = []
                progress = st.progress(0)
                for i, mail in enumerate(selected):
                    saved.append(save_company_email(mail, company or st.session_state.get("email_company_query", "AZIENDA"), archive_root=archive_root, skip_inline_logo=skip_inline))
                    progress.progress((i + 1) / max(1, len(selected)))
                st.session_state["email_last_save_results"] = saved
                st.success(f"Email selezionate elaborate: {len(saved)}")

    saved_results = st.session_state.get("email_last_save_results") or []
    if saved_results:
        st.divider()
        st.subheader("Risultato archiviazione email")
        rows = []
        for r in saved_results:
            rows.append({
                "ok": r.get("ok"),
                "duplicata": r.get("duplicate", False),
                "cartella": r.get("folder") or ((r.get("existing") or {}).get("pdf_path") if isinstance(r.get("existing"), dict) else ""),
                "pdf": r.get("pdf"),
                "allegati_salvati": len(r.get("attachments") or []),
                "duplicati_file": len(r.get("duplicates") or []),
                "messaggio": r.get("message"),
            })
        if pd:
            st.dataframe(pd.DataFrame(rows), width="stretch", hide_index=True)
        else:
            st.table(rows)
        last_folder = next((r.get("folder") for r in saved_results if r.get("folder")), "")
        if last_folder:
            st.code(last_folder)


def render_review_queue() -> None:
    render_header()
    if not require_perm("review"):
        st.error("Ruolo non autorizzato alla correzione documenti.")
        return
    st.subheader("Coda documenti da verificare")
    rows = query_all(
        """
        SELECT q.id qid, q.documento_id, q.file, q.categoria_suggerita, q.cliente_suggerito,
               q.motivo_incertezza, q.data_creazione, d.testo_estratto, d.metadata_json, d.nome_file_originale
        FROM coda_da_verificare q JOIN documenti d ON d.id=q.documento_id
        WHERE q.stato_lavorazione='aperta'
        ORDER BY q.id DESC
        """
    )
    if not rows:
        st.success("Nessun documento in coda.")
        return
    labels = [f"#{r['documento_id']} - {Path(r['file']).name} - {r['categoria_suggerita']} - {r['cliente_suggerito'] or 'cliente incerto'}" for r in rows]
    idx = st.selectbox("Documento", range(len(rows)), format_func=lambda i: labels[i])
    row = rows[idx]
    st.warning(f"Motivo incertezza: {row.get('motivo_incertezza')}")
    with st.expander("Testo estratto", expanded=True):
        st.text_area("OCR/testo", value=(row.get("testo_estratto") or "")[:20000], height=260)
    clients = query_all("SELECT id, ragione_sociale, partita_iva FROM clienti ORDER BY ragione_sociale")
    client_options = ["-- crea nuovo cliente --"] + [f"{c['id']} | {c['ragione_sociale']} | {c.get('partita_iva') or ''}" for c in clients]
    with st.form("review_form"):
        categoria = st.selectbox("Categoria corretta", DOCUMENT_CATEGORIES, index=max(0, DOCUMENT_CATEGORIES.index(row["categoria_suggerita"]) if row["categoria_suggerita"] in DOCUMENT_CATEGORIES else 0))
        client_choice = st.selectbox("Cliente corretto", client_options)
        st.markdown("**Nuovo cliente / dati da aggiornare**")
        ragione = st.text_input("Ragione sociale", value=row.get("cliente_suggerito") or "")
        piva = st.text_input("Partita IVA")
        cf = st.text_input("Codice fiscale")
        sede = st.text_input("Sede")
        admin = st.text_input("Amministratore")
        ateco = st.text_input("ATECO")
        email = st.text_input("Email")
        pec = st.text_input("PEC")
        notes = st.text_area("Note correzione")
        submitted = st.form_submit_button("Conferma archiviazione e apprendi regola", type="primary")
    if submitted:
        selected_client_id = None
        if client_choice != "-- crea nuovo cliente --":
            selected_client_id = int(client_choice.split("|", 1)[0].strip())
            cdata = query_one("SELECT * FROM clienti WHERE id=?", (selected_client_id,)) or {}
            ragione = ragione or cdata.get("ragione_sociale") or ""
            piva = piva or cdata.get("partita_iva") or ""
            cf = cf or cdata.get("codice_fiscale") or ""
        c = ClassificationResult(
            categoria=categoria,
            confidence=0.99,
            ragione_sociale=ragione,
            partita_iva=piva,
            codice_fiscale=cf,
            sede=sede,
            amministratore=admin,
            ateco=ateco,
            email=email,
            pec=pec,
            data_documento="",
            importo=None,
            metadata={"correzione_note": notes},
            motivazione="Correzione manuale utente",
            anomalie=[],
            needs_review=False,
            matched_client_id=selected_client_id,
        )
        client_id = selected_client_id or ensure_client(c, create=True)
        c.matched_client_id = client_id
        doc = query_one("SELECT * FROM documenti WHERE id=?", (row["documento_id"],))
        if not doc:
            st.error("Documento non trovato.")
            return
        old = Path(doc["percorso_file"])
        date_doc = doc.get("data_documento") or extract_document_date(doc.get("testo_estratto") or "")
        c.data_documento = date_doc
        c.importo = doc.get("importo")
        target = build_target_path(old, c, "archiviato")
        target.parent.mkdir(parents=True, exist_ok=True)
        target = unique_path(target)
        if old.exists():
            shutil.move(str(old), str(target))
        else:
            target = old
        meta = json.loads(doc.get("metadata_json") or "{}")
        meta["correzione_manual"] = asdict(c)
        con = connect()
        try:
            con.execute(
                """
                UPDATE documenti SET cliente_id=?, categoria=?, nome_file_archiviato=?, percorso_file=?, stato_riconoscimento='archiviato',
                    confidenza=?, metadata_json=?, data_ultimo_aggiornamento=? WHERE id=?
                """,
                (client_id, categoria, target.name, str(target), c.confidence, json.dumps(meta, ensure_ascii=False), now_iso(), row["documento_id"]),
            )
            con.execute("UPDATE coda_da_verificare SET stato_lavorazione='chiusa', data_chiusura=? WHERE id=?", (now_iso(), row["qid"]))
            try:
                con.execute("DELETE FROM documenti_fts WHERE documento_id=?", (row["documento_id"],))
                con.execute("INSERT INTO documenti_fts(documento_id,titolo,corpo,metadata) VALUES(?,?,?,?)", (row["documento_id"], target.name, doc.get("testo_estratto") or "", json.dumps(meta, ensure_ascii=False)))
            except Exception:
                pass
            con.commit()
        finally:
            con.close()
        learn_from_confirmation(categoria, ragione, piva, cf, doc.get("testo_estratto") or "", "correzione_manual")
        log_event("CORREZIONE_CODA", target.name, "ARCHIVIATO", notes)
        st.success("Documento corretto, archiviato e regola appresa.")
        st.rerun()


def search_documents(query: str, category: str, status: str) -> list[dict[str, Any]]:
    params: list[Any] = []
    base = """
        SELECT d.*, c.ragione_sociale, c.partita_iva, c.codice_fiscale
        FROM documenti d LEFT JOIN clienti c ON c.id=d.cliente_id
    """
    where = []
    if category and category != "Tutte":
        where.append("d.categoria=?")
        params.append(category)
    if status and status != "Tutti":
        where.append("d.stato_riconoscimento=?")
        params.append(status)
    q = normalize_spaces(query)
    if q:
        like = f"%{q}%"
        where.append("(d.nome_file_archiviato LIKE ? OR d.testo_estratto LIKE ? OR d.metadata_json LIKE ? OR c.ragione_sociale LIKE ? OR c.partita_iva LIKE ? OR c.codice_fiscale LIKE ?)")
        params.extend([like, like, like, like, like, like])
    sql = base + (" WHERE " + " AND ".join(where) if where else "") + " ORDER BY d.id DESC LIMIT 500"
    return query_all(sql, tuple(params))


def render_search() -> None:
    render_header()
    if not require_perm("read"):
        st.error("Ruolo non autorizzato alla consultazione.")
        return
    st.subheader("Ricerca full-text archivio")
    c1, c2, c3 = st.columns([2, 1, 1])
    q = c1.text_input("Cerca per ragione sociale, P.IVA, CF, amministratore, banca, importo, IBAN, ATECO, protocollo o testo")
    category = c2.selectbox("Categoria", ["Tutte"] + DOCUMENT_CATEGORIES)
    status = c3.selectbox("Stato", ["Tutti", "archiviato", "da_verificare"])
    rows = search_documents(q, category, status)
    st.caption(f"Risultati: {len(rows)}")
    if rows and pd:
        df = pd.DataFrame(rows)
        cols = [c for c in ["id", "ragione_sociale", "partita_iva", "categoria", "nome_file_archiviato", "data_documento", "importo", "stato_riconoscimento", "confidenza", "percorso_file"] if c in df.columns]
        st.dataframe(df[cols], width="stretch", hide_index=True, height=380)
    elif rows:
        st.json(rows[:100])
    if rows:
        ids = [int(r["id"]) for r in rows]
        selected = st.selectbox("Apri risultato", ids)
        doc = next((r for r in rows if int(r["id"]) == selected), None)
        if doc:
            path = Path(doc["percorso_file"])
            st.write(f"**File:** `{path}`")
            st.write(f"**Anteprima testo:** {(doc.get('testo_estratto') or '')[:800]}")
            if path.exists():
                with path.open("rb") as f:
                    st.download_button("Scarica documento", data=f.read(), file_name=path.name)
            else:
                st.warning("File non trovato sul disco. Il record e' presente nel database.")


def render_clients() -> None:
    render_header()
    st.subheader("Anagrafica clienti")
    rows = query_all("SELECT * FROM clienti ORDER BY ragione_sociale")
    if rows and pd:
        st.dataframe(pd.DataFrame(rows), width="stretch", hide_index=True, height=360)
    elif rows:
        st.table(rows)
    else:
        st.info("Nessun cliente censito.")
    if require_perm("review"):
        with st.expander("Crea / aggiorna cliente manualmente"):
            with st.form("client_form"):
                ragione = st.text_input("Ragione sociale")
                piva = st.text_input("Partita IVA")
                cf = st.text_input("Codice fiscale")
                sede = st.text_input("Sede")
                admin = st.text_input("Amministratore")
                ateco = st.text_input("ATECO")
                email = st.text_input("Email")
                pec = st.text_input("PEC")
                ok = st.form_submit_button("Salva cliente")
            if ok and ragione:
                c = ClassificationResult("Documento generico", 0.0, ragione, piva, cf, sede, admin, ateco, email, pec)
                cid = ensure_client(c, create=True)
                log_event("CLIENTE_MANUALE", ragione, "OK", f"ID {cid}")
                st.success("Cliente salvato.")
                st.rerun()


def render_reports_exports() -> None:
    render_header()
    if not require_perm("export"):
        st.error("Ruolo non autorizzato a report/export.")
        return
    st.subheader("Report PDF ed esportazioni")
    c1, c2 = st.columns(2)
    with c1:
        st.markdown("### Report PDF")
        if st.button("Genera report PDF generale", type="primary"):
            path = generate_pdf_report("generale")
            st.success(f"Report creato: {path.name}")
            with path.open("rb") as f:
                st.download_button("Scarica report", f.read(), file_name=path.name)
    with c2:
        st.markdown("### Export dati")
        table = st.selectbox("Tabella", ["clienti", "documenti", "log_attivita", "coda_da_verificare", "modello_apprendimento"])
        csv_data = export_table_csv(table)
        st.download_button("Scarica CSV", csv_data, file_name=f"{table}.csv", mime="text/csv")
        excel_data = export_table_excel(table)
        st.download_button("Scarica Excel", excel_data, file_name=f"{table}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    st.divider()
    st.subheader("Report disponibili")
    files = sorted(REPORT_DIR.glob("*"), key=lambda x: x.stat().st_mtime, reverse=True)
    for f in files[:20]:
        st.write(f"`{f.name}` - {datetime.fromtimestamp(f.stat().st_mtime).strftime('%d-%m-%Y %H:%M')}")


def render_backup_integrations() -> None:
    render_header()
    if not require_perm("backup"):
        st.error("Ruolo non autorizzato a backup/impostazioni cloud.")
        return
    st.subheader("Backup locale, cloud e integrazioni")
    st.markdown("### Backup")
    include_temp = st.checkbox("Includi cartella temporanea nel backup", value=False)
    external_folder = st.text_input("Cartella esterna opzionale per copia backup", placeholder=r"Esempio: D:\Backup_FinancePlus")
    if st.button("Crea backup ZIP", type="primary"):
        z = create_backup_zip(include_temp)
        copied = ""
        if external_folder:
            dest_dir = Path(external_folder).expanduser()
            dest_dir.mkdir(parents=True, exist_ok=True)
            copied_path = dest_dir / z.name
            shutil.copy2(z, copied_path)
            copied = f" Copiato anche in {copied_path}."
        st.success(f"Backup creato: {z}.{copied}")
        with z.open("rb") as f:
            st.download_button("Scarica backup", f.read(), file_name=z.name, mime="application/zip")
    st.divider()
    st.markdown("### Predisposizione Google Drive / pCloud / Gmail / PEC / scanner")
    with st.form("integrations"):
        google_drive = st.text_input("Cartella Google Drive locale o mount rclone", value=get_setting("google_drive_path"))
        pcloud = st.text_input("Cartella pCloud locale/WebDAV mount", value=get_setting("pcloud_path"))
        pec = st.text_input("PEC/IMAP host", value=get_setting("pec_host"))
        gmail = st.text_input("Gmail label o query import futura", value=get_setting("gmail_query"))
        scanner = st.text_input("Cartella scanner di rete", value=get_setting("scanner_folder"))
        api_enabled = st.checkbox("Predisponi API/integrator esterni", value=get_setting("api_enabled", "0") == "1")
        ok = st.form_submit_button("Salva configurazione")
    if ok:
        set_setting("google_drive_path", google_drive)
        set_setting("pcloud_path", pcloud)
        set_setting("pec_host", pec)
        set_setting("gmail_query", gmail)
        set_setting("scanner_folder", scanner)
        set_setting("api_enabled", "1" if api_enabled else "0")
        log_event("IMPOSTAZIONI_INTEGRAZIONI", "", "OK")
        st.success("Configurazione salvata. Gli adapter cloud/API sono predisposti ma non inviano dati senza credenziali esplicite.")


def render_users_settings() -> None:
    render_header()
    if not require_perm("users"):
        st.error("Solo Amministratore puo' gestire utenti e permessi.")
        return
    st.subheader("Permessi utenti")
    users = query_all("SELECT id, username, role, active, created_at, updated_at FROM users ORDER BY username")
    if pd:
        st.dataframe(pd.DataFrame(users), width="stretch", hide_index=True)
    else:
        st.table(users)
    with st.expander("Crea nuovo utente"):
        with st.form("new_user"):
            username = st.text_input("Username")
            password = st.text_input("Password", type="password")
            role = st.selectbox("Ruolo", list(ROLE_PERMS.keys()))
            ok = st.form_submit_button("Crea utente")
        if ok and username and password:
            try:
                execute(
                    "INSERT INTO users(username,password_hash,role,active,created_at,updated_at) VALUES(?,?,?,?,?,?)",
                    (username.strip(), hash_password(password), role, 1, now_iso(), now_iso()),
                )
                log_event("UTENTE_CREATO", username, "OK", role)
                st.success("Utente creato.")
                st.rerun()
            except sqlite3.IntegrityError:
                st.error("Username gia' esistente.")
    with st.expander("Cambio password utente corrente"):
        with st.form("change_password"):
            new_pw = st.text_input("Nuova password", type="password")
            okpw = st.form_submit_button("Aggiorna password")
        if okpw and new_pw:
            user = st.session_state.get("user") or {}
            execute("UPDATE users SET password_hash=?, updated_at=? WHERE id=?", (hash_password(new_pw), now_iso(), user.get("id")))
            log_event("PASSWORD_CAMBIATA", user.get("username", ""), "OK")
            st.success("Password aggiornata.")


def render_architecture() -> None:
    render_header()
    st.subheader("Specifiche integrate nel progetto")
    st.markdown(
        """
        Questa versione integra la catena IDP completa:

        1. **Ingestion**: upload file, scansione cartelle/sottocartelle, email IMAP con CERCA AZIENDA, Gmail/PEC/scanner/cloud.
        2. **Hash SHA256**: blocco duplicati prima dell'archiviazione.
        3. **OCR/lettura**: PyMuPDF per PDF digitali, pytesseract per immagini e scansioni, predisposizione OCR cloud/HTR.
        4. **Classificazione**: regole testuali, pattern fiscali, parole chiave, modello IA opzionale e apprendimento.
        5. **Estrazione metadati**: dati generali e campi specifici per categoria.
        6. **Riconoscimento cliente**: priorita a P.IVA, CF, ragione sociale, amministratore, sede, email/PEC e anagrafica.
        7. **Archiviazione automatica**: cliente / anno / mese / tipologia documentale.
        8. **Coda da verificare**: gestione manuale dei casi incerti con apprendimento della correzione.
        9. **Ricerca full-text**: testo OCR, metadati, P.IVA, CF, IBAN, protocollo, ATECO, importi.
        10. **Report e export**: PDF, CSV, Excel.
        11. **Backup**: database, archivio, log, configurazioni e modello appreso.
        12. **Permessi**: Amministratore, Operatore, Solo lettura.
        """
    )
    st.code(
        """Archivio_Documenti/
    Cliente_RagioneSociale_PIVA/
        2026/
            Luglio/
                Visura/
                Bilancio/
                Centrale_Rischi/
                DURC/
                Contratti/
                Fatture/
                Estratti_Conto/
                Preventivi/
                Business_Plan/
                Report_Bancari/
                Dichiarazioni_Fiscali/
                Altro/
    Da_Verificare/""",
        language="text",
    )
    st.subheader("Funzioni obbligatorie")
    funzioni = [
        ["OCR su PDF e immagini", "Si"],
        ["Lettura documenti scannerizzati", "Si"],
        ["Classificazione automatica", "Si"],
        ["Estrazione metadati", "Si"],
        ["Riconoscimento cliente da P.IVA, CF o ragione sociale", "Si"],
        ["Rinomina automatica file", "Si"],
        ["SQLite locale", "Si"],
        ["Predisposizione PostgreSQL/API/cloud", "Si"],
        ["Coda da verificare", "Si"],
        ["Apprendimento continuo", "Si"],
    ]
    if pd:
        st.dataframe(pd.DataFrame(funzioni, columns=["Funzione", "Stato"]), width="stretch", hide_index=True)
    else:
        st.table(funzioni)


def main() -> None:
    if st is None:
        print("Streamlit non installato. Esegui: pip install streamlit")
        raise SystemExit(1)
    st.set_page_config(page_title=APP_TITLE, page_icon="FP", layout="wide")
    init_db()
    if "user" not in st.session_state:
        render_login()
        return
    with st.sidebar:
        logo_path = ROOT_DIR / "assets" / "financeplus_archive_ai_idp_logo.png"
        if logo_path.exists():
            st.image(str(logo_path), width="stretch")
        else:
            st.markdown("## FinancePlus")
        st.caption(f"DB: {DB_PATH.name}")
        pages = [
            "Dashboard",
            "Importa e cataloga",
            "Email azienda",
            "Coda da verificare",
            "Ricerca archivio",
            "Clienti",
            "Report ed export",
            "Backup e integrazioni",
            "Utenti e permessi",
            "Architettura integrata",
        ]
        page = st.radio("Menu", pages)
        st.divider()
        if st.button("Logout"):
            log_event("LOGOUT", st.session_state.get("user", {}).get("username", ""), "OK")
            del st.session_state["user"]
            st.rerun()
    if page == "Dashboard":
        render_dashboard()
    elif page == "Importa e cataloga":
        render_import()
    elif page == "Email azienda":
        render_email_company()
    elif page == "Coda da verificare":
        render_review_queue()
    elif page == "Ricerca archivio":
        render_search()
    elif page == "Clienti":
        render_clients()
    elif page == "Report ed export":
        render_reports_exports()
    elif page == "Backup e integrazioni":
        render_backup_integrations()
    elif page == "Utenti e permessi":
        render_users_settings()
    elif page == "Architettura integrata":
        render_architecture()


if __name__ == "__main__":
    main()
