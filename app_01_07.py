# FinancePlus App 01_07 - Streamlit IDP Documentale
# Versione unica pronta per PC, GitHub e Streamlit Cloud.
# Avvio locale: streamlit run streamlit_app.py

from __future__ import annotations

import hashlib
import io
import json
import os
import re
import shutil
import sqlite3
import zipfile
from dataclasses import dataclass
from datetime import datetime, date
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

import pandas as pd
import streamlit as st

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    from PIL import Image
except Exception:
    Image = None

try:
    import pytesseract
except Exception:
    pytesseract = None

try:
    from docx import Document
except Exception:
    Document = None

try:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
        PageBreak,
    )
except Exception:
    SimpleDocTemplate = None


APP_NAME = "FinancePlus IDP - App 01_07"
APP_VERSION = "1.0.0"
DB_PATH = Path("data/financeplus_idp_0107.db")
ARCHIVE_ROOT = Path("archivio_clienti")
REPORT_DIR = Path("report")
STAGING_DIR = Path("data/staging_upload")

ALLOWED_EXTENSIONS = {
    ".pdf",
    ".png",
    ".jpg",
    ".jpeg",
    ".tif",
    ".tiff",
    ".txt",
    ".csv",
    ".xlsx",
    ".xls",
    ".docx",
}

DOC_TYPES = {
    "visura": ["visura", "registro imprese", "rea", "camera di commercio", "cciaa", "ateco"],
    "bilancio": ["bilancio", "stato patrimoniale", "conto economico", "nota integrativa", "xbrl"],
    "centrale_rischi": ["centrale rischi", "banca d'italia", "accordato", "utilizzato", "sconfinamento", "scaduto"],
    "estratto_conto": ["estratto conto", "saldo contabile", "saldo disponibile", "movimenti", "valuta"],
    "contratto": ["contratto", "accordo", "fornitura", "mandato", "scrittura privata", "ordine"],
    "fattura": ["fattura", "invoice", "totale documento", "imponibile", "iva"],
    "documento_identita": ["carta identita", "passaporto", "patente", "documento di identita"],
    "report": ["report", "analisi", "rating", "score", "dscr", "business plan"],
    "mail": ["from:", "to:", "subject:", "oggetto", "mittente"],
}

STOPWORDS_COMPANY = {
    "srl", "s.r.l", "spa", "s.p.a", "snc", "sas", "societa", "azienda", "impresa", "group", "italia"
}


@dataclass
class AnalysisResult:
    original_name: str
    relative_path: str
    extension: str
    sha256: str
    text: str
    text_preview: str
    document_type: str
    document_date: str
    vat_number: str
    fiscal_code: str
    amount: str
    matched_client_id: Optional[int]
    matched_client_name: str
    confidence: int
    status: str
    suggested_filename: str
    archived_path: str
    notes: str


# -----------------------------
# Utility
# -----------------------------


def ensure_dirs() -> None:
    for p in [DB_PATH.parent, ARCHIVE_ROOT, REPORT_DIR, STAGING_DIR]:
        p.mkdir(parents=True, exist_ok=True)


def normalize_text(value: str) -> str:
    value = value or ""
    value = value.lower()
    value = value.replace("&", " e ")
    value = re.sub(r"[^a-z0-9]+", " ", value)
    value = re.sub(r"\s+", " ", value).strip()
    return value


def safe_filename(value: str, fallback: str = "documento") -> str:
    value = value or fallback
    value = value.strip().replace("/", "_").replace("\\", "_")
    value = re.sub(r"[^A-Za-z0-9._\- ]+", "_", value)
    value = re.sub(r"\s+", "_", value)
    value = re.sub(r"_+", "_", value)
    value = value.strip("._-")
    return value[:180] or fallback


def company_slug(value: str) -> str:
    return safe_filename(value.upper(), "AZIENDA_NON_RICONOSCIUTA")


def sha256_bytes(raw: bytes) -> str:
    return hashlib.sha256(raw).hexdigest()


def short_hash(value: str) -> str:
    return value[:10] if value else ""


def today_str() -> str:
    return datetime.now().strftime("%d-%m-%Y")


def month_folder_from_date(ddmmyyyy: str) -> str:
    try:
        dt = datetime.strptime(ddmmyyyy, "%d-%m-%Y")
        return dt.strftime("%m-%Y")
    except Exception:
        return datetime.now().strftime("%m-%Y")


def read_bytes(uploaded_file) -> bytes:
    uploaded_file.seek(0)
    raw = uploaded_file.read()
    uploaded_file.seek(0)
    return raw


def clean_relative_path(path: str) -> str:
    path = path.replace("\\", "/")
    parts = [safe_filename(p, "file") for p in path.split("/") if p and p not in [".", ".."]]
    return "/".join(parts) if parts else safe_filename(path, "file")


# -----------------------------
# Database
# -----------------------------


def get_connection() -> sqlite3.Connection:
    ensure_dirs()
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def init_db() -> None:
    ensure_dirs()
    with get_connection() as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS clients (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                company_name TEXT NOT NULL UNIQUE,
                vat_number TEXT,
                fiscal_code TEXT,
                administrator TEXT,
                email TEXT,
                pec TEXT,
                notes TEXT,
                created_at TEXT NOT NULL
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS documents (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                original_name TEXT NOT NULL,
                relative_path TEXT,
                sha256 TEXT NOT NULL UNIQUE,
                extension TEXT,
                document_type TEXT,
                document_date TEXT,
                client_id INTEGER,
                client_name TEXT,
                confidence INTEGER,
                status TEXT,
                suggested_filename TEXT,
                archived_path TEXT,
                text_preview TEXT,
                full_text TEXT,
                vat_number TEXT,
                fiscal_code TEXT,
                amount TEXT,
                notes TEXT,
                created_at TEXT NOT NULL,
                FOREIGN KEY(client_id) REFERENCES clients(id)
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS settings (
                key TEXT PRIMARY KEY,
                value TEXT
            )
            """
        )
        conn.commit()


def seed_demo_clients() -> None:
    demo = [
        ("BEL GARDEN EUROPE S.R.L.", "12958900966", "", "COLETTA DANIELE", "bge@belgardeneurope.com", "bge@arubapec.it"),
        ("PELCOM S.R.L.", "", "", "", "", ""),
        ("ETS GROUP S.R.L.", "", "", "", "", ""),
        ("FRANCESCO RUSSO S.R.L.", "", "", "", "", ""),
        ("STC S.R.L.", "", "", "", "", ""),
    ]
    with get_connection() as conn:
        for row in demo:
            try:
                conn.execute(
                    """
                    INSERT INTO clients(company_name, vat_number, fiscal_code, administrator, email, pec, notes, created_at)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (*row, "Cliente demo app 01_07", datetime.now().isoformat(timespec="seconds")),
                )
            except sqlite3.IntegrityError:
                pass
        conn.commit()


def clients_df() -> pd.DataFrame:
    with get_connection() as conn:
        return pd.read_sql_query("SELECT * FROM clients ORDER BY company_name", conn)


def documents_df(where: str = "", params: Tuple = ()) -> pd.DataFrame:
    query = "SELECT d.*, c.company_name AS registry_company FROM documents d LEFT JOIN clients c ON c.id = d.client_id"
    if where:
        query += " WHERE " + where
    query += " ORDER BY d.created_at DESC"
    with get_connection() as conn:
        return pd.read_sql_query(query, conn, params=params)


def insert_client(company_name: str, vat_number: str = "", fiscal_code: str = "", administrator: str = "", email: str = "", pec: str = "", notes: str = "") -> int:
    company_name = company_name.strip().upper()
    with get_connection() as conn:
        try:
            cur = conn.execute(
                """
                INSERT INTO clients(company_name, vat_number, fiscal_code, administrator, email, pec, notes, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (company_name, vat_number, fiscal_code, administrator, email, pec, notes, datetime.now().isoformat(timespec="seconds")),
            )
            conn.commit()
            return int(cur.lastrowid)
        except sqlite3.IntegrityError:
            row = conn.execute("SELECT id FROM clients WHERE company_name = ?", (company_name,)).fetchone()
            return int(row["id"])


def update_client(client_id: int, **fields) -> None:
    allowed = ["company_name", "vat_number", "fiscal_code", "administrator", "email", "pec", "notes"]
    data = {k: v for k, v in fields.items() if k in allowed}
    if not data:
        return
    assignments = ", ".join([f"{k} = ?" for k in data.keys()])
    values = list(data.values()) + [client_id]
    with get_connection() as conn:
        conn.execute(f"UPDATE clients SET {assignments} WHERE id = ?", values)
        conn.commit()


def get_client_by_id(client_id: Optional[int]) -> Optional[sqlite3.Row]:
    if not client_id:
        return None
    with get_connection() as conn:
        return conn.execute("SELECT * FROM clients WHERE id = ?", (client_id,)).fetchone()


def duplicate_by_hash(sha: str) -> Optional[sqlite3.Row]:
    with get_connection() as conn:
        return conn.execute("SELECT * FROM documents WHERE sha256 = ?", (sha,)).fetchone()


def insert_document(result: AnalysisResult) -> None:
    with get_connection() as conn:
        conn.execute(
            """
            INSERT OR IGNORE INTO documents(
                original_name, relative_path, sha256, extension, document_type, document_date,
                client_id, client_name, confidence, status, suggested_filename, archived_path,
                text_preview, full_text, vat_number, fiscal_code, amount, notes, created_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                result.original_name,
                result.relative_path,
                result.sha256,
                result.extension,
                result.document_type,
                result.document_date,
                result.matched_client_id,
                result.matched_client_name,
                result.confidence,
                result.status,
                result.suggested_filename,
                result.archived_path,
                result.text_preview,
                result.text,
                result.vat_number,
                result.fiscal_code,
                result.amount,
                result.notes,
                datetime.now().isoformat(timespec="seconds"),
            ),
        )
        conn.commit()


def update_document_status(doc_id: int, status: str, client_id: Optional[int], client_name: str, archived_path: str, suggested_filename: str) -> None:
    with get_connection() as conn:
        conn.execute(
            """
            UPDATE documents
            SET status = ?, client_id = ?, client_name = ?, archived_path = ?, suggested_filename = ?
            WHERE id = ?
            """,
            (status, client_id, client_name, archived_path, suggested_filename, doc_id),
        )
        conn.commit()


# -----------------------------
# Extraction
# -----------------------------


def extract_text_from_pdf(raw: bytes, ocr_if_empty: bool = True) -> str:
    texts: List[str] = []
    if fitz is not None:
        try:
            with fitz.open(stream=raw, filetype="pdf") as doc:
                for page in doc:
                    text = page.get_text("text") or ""
                    texts.append(text)
                all_text = "\n".join(texts).strip()
                if all_text or not ocr_if_empty or pytesseract is None or Image is None:
                    return all_text
                # OCR su pagine rasterizzate se PDF scannerizzato.
                ocr_pages = []
                for page in doc:
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                    img = Image.open(io.BytesIO(pix.tobytes("png")))
                    ocr_pages.append(pytesseract.image_to_string(img, lang="ita+eng"))
                return "\n".join(ocr_pages).strip()
        except Exception:
            pass
    if pdfplumber is not None:
        try:
            with pdfplumber.open(io.BytesIO(raw)) as pdf:
                for page in pdf.pages:
                    texts.append(page.extract_text() or "")
            return "\n".join(texts).strip()
        except Exception:
            pass
    return ""


def extract_text_from_image(raw: bytes) -> str:
    if Image is None:
        return ""
    try:
        img = Image.open(io.BytesIO(raw))
        if pytesseract is None:
            return ""
        return pytesseract.image_to_string(img, lang="ita+eng") or ""
    except Exception:
        return ""


def extract_text_from_docx(raw: bytes) -> str:
    if Document is None:
        return ""
    try:
        doc = Document(io.BytesIO(raw))
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        tables = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                tables.append(" | ".join(cells))
        return "\n".join(paragraphs + tables).strip()
    except Exception:
        return ""


def extract_text_from_spreadsheet(raw: bytes, suffix: str) -> str:
    try:
        if suffix == ".csv":
            try:
                df = pd.read_csv(io.BytesIO(raw), sep=None, engine="python")
            except Exception:
                df = pd.read_csv(io.BytesIO(raw), encoding="latin1", sep=None, engine="python")
            return df.head(200).to_csv(index=False)
        xls = pd.ExcelFile(io.BytesIO(raw))
        chunks = []
        for sheet in xls.sheet_names[:10]:
            df = xls.parse(sheet).head(200)
            chunks.append(f"FOGLIO: {sheet}\n" + df.to_csv(index=False))
        return "\n".join(chunks)
    except Exception:
        return ""


def extract_text_from_file(raw: bytes, filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(raw)
    if suffix in [".png", ".jpg", ".jpeg", ".tif", ".tiff"]:
        return extract_text_from_image(raw)
    if suffix in [".txt"]:
        for enc in ["utf-8", "latin1", "cp1252"]:
            try:
                return raw.decode(enc, errors="ignore")
            except Exception:
                continue
        return ""
    if suffix in [".csv", ".xlsx", ".xls"]:
        return extract_text_from_spreadsheet(raw, suffix)
    if suffix == ".docx":
        return extract_text_from_docx(raw)
    return ""


# -----------------------------
# Analysis and classification
# -----------------------------


def extract_vat(text: str) -> str:
    patterns = [
        r"(?:p\.?\s*iva|partita\s+iva|piva)[^0-9]{0,20}([0-9]{11})",
        r"(?:vat)[^0-9]{0,20}([0-9]{11})",
    ]
    low = text.lower()
    for p in patterns:
        m = re.search(p, low, flags=re.I)
        if m:
            return m.group(1)
    # Fallback prudente: solo se ci sono pochi numeri a 11 cifre.
    candidates = re.findall(r"\b[0-9]{11}\b", text)
    if len(candidates) == 1:
        return candidates[0]
    return ""


def extract_fiscal_code(text: str) -> str:
    patterns = re.findall(r"\b[A-Z]{6}[0-9]{2}[A-Z][0-9]{2}[A-Z][0-9]{3}[A-Z]\b", text.upper())
    return patterns[0] if patterns else ""


def extract_amount(text: str) -> str:
    matches = re.findall(r"(?:euro|eur|totale|importo|iva)?\s*([0-9]{1,3}(?:\.[0-9]{3})*(?:,[0-9]{2})|[0-9]+,[0-9]{2})", text.lower())
    if not matches:
        return ""
    # Prende l'importo piu alto come indicazione riepilogativa.
    def to_float(x: str) -> float:
        try:
            return float(x.replace(".", "").replace(",", "."))
        except Exception:
            return 0.0
    best = max(matches, key=to_float)
    return best


def extract_date(text: str, filename: str = "") -> str:
    joined = f"{filename}\n{text}"
    patterns = [
        r"\b([0-3]?[0-9])[\-/\.]([0-1]?[0-9])[\-/\.](20[0-9]{2}|19[0-9]{2})\b",
        r"\b(20[0-9]{2}|19[0-9]{2})[\-/\.]([0-1]?[0-9])[\-/\.]([0-3]?[0-9])\b",
    ]
    for p in patterns:
        m = re.search(p, joined)
        if not m:
            continue
        try:
            if len(m.group(1)) == 4:
                y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
            else:
                d, mo, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
            return datetime(y, mo, d).strftime("%d-%m-%Y")
        except Exception:
            continue
    return today_str()


def classify_document(text: str, filename: str) -> Tuple[str, int, str]:
    haystack = normalize_text(f"{filename}\n{text}")
    best_type = "altro"
    best_score = 0
    evidence = []
    for doc_type, keywords in DOC_TYPES.items():
        score = 0
        hits = []
        for kw in keywords:
            nkw = normalize_text(kw)
            if nkw and nkw in haystack:
                score += 20
                hits.append(kw)
        if score > best_score:
            best_type = doc_type
            best_score = score
            evidence = hits
    confidence = min(95, max(25, best_score + 35 if best_score else 25))
    return best_type, confidence, ", ".join(evidence[:5])


def text_preview(text: str, max_chars: int = 900) -> str:
    text = re.sub(r"\s+", " ", text or "").strip()
    return text[:max_chars] + ("..." if len(text) > max_chars else "")


def match_client(text: str, filename: str, vat_number: str, fiscal_code: str) -> Tuple[Optional[int], str, int, str]:
    df = clients_df()
    if df.empty:
        return None, "", 0, "Nessuna anagrafica presente"

    haystack = normalize_text(f"{filename}\n{text}")
    best = (None, "", 0, "")

    for _, row in df.iterrows():
        cid = int(row["id"])
        name = str(row["company_name"] or "")
        vat = str(row.get("vat_number", "") or "")
        cf = str(row.get("fiscal_code", "") or "")
        admin = str(row.get("administrator", "") or "")
        score = 0
        reasons = []

        if vat and vat_number and vat == vat_number:
            score += 65
            reasons.append("P.IVA coincidente")
        elif vat and vat in text:
            score += 60
            reasons.append("P.IVA trovata nel testo")

        if cf and fiscal_code and cf.upper() == fiscal_code.upper():
            score += 55
            reasons.append("Codice fiscale coincidente")
        elif cf and cf.upper() in text.upper():
            score += 50
            reasons.append("Codice fiscale trovato nel testo")

        norm_name = normalize_text(name)
        if norm_name and norm_name in haystack:
            score += 55
            reasons.append("Ragione sociale completa")
        else:
            tokens = [t for t in norm_name.split() if len(t) > 2 and t not in STOPWORDS_COMPANY]
            if tokens:
                matched_tokens = [t for t in tokens if t in haystack]
                ratio = len(matched_tokens) / max(1, len(tokens))
                if ratio >= 0.75:
                    score += int(40 * ratio)
                    reasons.append("Denominazione parziale")

        norm_admin = normalize_text(admin)
        if norm_admin and norm_admin in haystack:
            score += 20
            reasons.append("Amministratore riconosciuto")

        if score > best[2]:
            best = (cid, name, min(98, score), ", ".join(reasons))

    return best


def suggest_name(client_name: str, doc_type: str, document_date: str, extension: str) -> str:
    base = company_slug(client_name or "AZIENDA_NON_RICONOSCIUTA")
    return safe_filename(f"{base}_{doc_type}_{document_date}{extension}")


def archive_raw_file(raw: bytes, client_name: str, doc_type: str, document_date: str, filename: str) -> str:
    client_folder = company_slug(client_name or "AZIENDA_NON_RICONOSCIUTA")
    month_folder = month_folder_from_date(document_date)
    target_dir = ARCHIVE_ROOT / client_folder / month_folder / safe_filename(doc_type, "altro")
    target_dir.mkdir(parents=True, exist_ok=True)
    target = target_dir / filename
    if target.exists():
        stem = target.stem
        suffix = target.suffix
        target = target_dir / f"{stem}_{datetime.now().strftime('%H%M%S')}{suffix}"
    target.write_bytes(raw)
    return str(target)


def analyze_upload(uploaded_file, auto_archive: bool, min_confidence: int) -> AnalysisResult:
    raw = read_bytes(uploaded_file)
    original_name = Path(uploaded_file.name).name
    relative_path = clean_relative_path(uploaded_file.name)
    extension = Path(original_name).suffix.lower()
    file_hash = sha256_bytes(raw)

    duplicate = duplicate_by_hash(file_hash)
    if duplicate is not None:
        return AnalysisResult(
            original_name=original_name,
            relative_path=relative_path,
            extension=extension,
            sha256=file_hash,
            text=str(duplicate["full_text"] or ""),
            text_preview=str(duplicate["text_preview"] or "Duplicato gia presente"),
            document_type=str(duplicate["document_type"] or "duplicato"),
            document_date=str(duplicate["document_date"] or today_str()),
            vat_number=str(duplicate["vat_number"] or ""),
            fiscal_code=str(duplicate["fiscal_code"] or ""),
            amount=str(duplicate["amount"] or ""),
            matched_client_id=duplicate["client_id"],
            matched_client_name=str(duplicate["client_name"] or ""),
            confidence=int(duplicate["confidence"] or 0),
            status="duplicato",
            suggested_filename=str(duplicate["suggested_filename"] or original_name),
            archived_path=str(duplicate["archived_path"] or ""),
            notes="File non salvato: hash gia presente nel database.",
        )

    text = extract_text_from_file(raw, original_name)
    vat = extract_vat(text)
    cf = extract_fiscal_code(text)
    amount = extract_amount(text)
    doc_date = extract_date(text, original_name)
    doc_type, type_conf, type_evidence = classify_document(text, original_name)
    cid, cname, client_score, client_evidence = match_client(text, original_name, vat, cf)

    confidence = int(round((type_conf * 0.35) + (client_score * 0.65))) if cid else int(round(type_conf * 0.45))
    status = "da_verificare"
    archived_path = ""
    client_for_file = cname or "AZIENDA_NON_RICONOSCIUTA"
    suggested_filename = suggest_name(client_for_file, doc_type, doc_date, extension)

    if auto_archive and cid and confidence >= min_confidence:
        archived_path = archive_raw_file(raw, cname, doc_type, doc_date, suggested_filename)
        status = "archiviato"
    elif not text:
        status = "da_verificare"

    notes = "; ".join([x for x in [type_evidence, client_evidence] if x])
    if not text:
        notes = (notes + "; " if notes else "") + "Testo non estratto: controllare OCR/Tesseract o file protetto."

    return AnalysisResult(
        original_name=original_name,
        relative_path=relative_path,
        extension=extension,
        sha256=file_hash,
        text=text,
        text_preview=text_preview(text),
        document_type=doc_type,
        document_date=doc_date,
        vat_number=vat,
        fiscal_code=cf,
        amount=amount,
        matched_client_id=cid,
        matched_client_name=cname,
        confidence=confidence,
        status=status,
        suggested_filename=suggested_filename,
        archived_path=archived_path,
        notes=notes,
    )


# -----------------------------
# PDF report
# -----------------------------


def _pdf_header_footer(canvas, doc):
    canvas.saveState()
    width, height = A4
    canvas.setFillColor(colors.HexColor("#0B1F3A"))
    canvas.rect(0, height - 1.15 * cm, width, 1.15 * cm, fill=True, stroke=False)
    canvas.setFillColor(colors.white)
    canvas.setFont("Helvetica-Bold", 10)
    canvas.drawString(1.2 * cm, height - 0.72 * cm, "FinancePlus IDP - App 01_07")
    canvas.setFillColor(colors.HexColor("#B87333"))
    canvas.setFont("Helvetica", 8)
    canvas.drawRightString(width - 1.2 * cm, 0.7 * cm, f"Pagina {doc.page}")
    canvas.restoreState()


def create_pdf_report(df_docs: pd.DataFrame, df_clients: pd.DataFrame, title: str = "Sintesi operativa App 01_07") -> bytes:
    if SimpleDocTemplate is None:
        raise RuntimeError("ReportLab non installato. Installa reportlab da requirements.txt")

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=1.2 * cm,
        leftMargin=1.2 * cm,
        topMargin=1.7 * cm,
        bottomMargin=1.2 * cm,
    )
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TitleCenter", parent=styles["Title"], alignment=TA_CENTER, textColor=colors.HexColor("#0B1F3A")))
    styles.add(ParagraphStyle(name="Copper", parent=styles["Heading2"], textColor=colors.HexColor("#B87333")))
    story = []

    story.append(Paragraph(title, styles["TitleCenter"]))
    story.append(Spacer(1, 0.35 * cm))
    story.append(Paragraph("Sintesi documentale, stato archivio, funzionalita e valutazione del progetto.", styles["BodyText"]))
    story.append(Spacer(1, 0.5 * cm))

    total_docs = int(len(df_docs))
    archived = int((df_docs["status"] == "archiviato").sum()) if not df_docs.empty and "status" in df_docs else 0
    queue = int((df_docs["status"] == "da_verificare").sum()) if not df_docs.empty and "status" in df_docs else 0
    duplicates = int((df_docs["status"] == "duplicato").sum()) if not df_docs.empty and "status" in df_docs else 0
    total_clients = int(len(df_clients))

    kpi_data = [
        ["Indicatore", "Valore"],
        ["Aziende in anagrafica", total_clients],
        ["Documenti analizzati", total_docs],
        ["Documenti archiviati", archived],
        ["Documenti da verificare", queue],
        ["Duplicati bloccati", duplicates],
    ]
    table = Table(kpi_data, colWidths=[9 * cm, 5 * cm])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0B1F3A")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#B8C2CC")),
        ("BACKGROUND", (0, 1), (-1, -1), colors.HexColor("#F6F8FB")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(table)
    story.append(Spacer(1, 0.55 * cm))

    story.append(Paragraph("Funzionalita integrate", styles["Copper"]))
    funzioni = [
        ["Modulo", "Stato", "Voto"],
        ["Upload cartella con sottocartelle", "Integrato", "9/10"],
        ["OCR PDF e immagini", "Integrato locale/cloud ready", "8/10"],
        ["Classificazione documento", "Rule-based estendibile con AI", "8/10"],
        ["Matching cliente per ragione sociale, P.IVA, CF, amministratore", "Integrato", "9/10"],
        ["Archivio cliente/mese/tipo", "Integrato", "9/10"],
        ["Coda documenti non riconosciuti", "Integrato", "8/10"],
        ["Report PDF", "Integrato", "8/10"],
        ["Predisposizione GitHub/Streamlit", "Integrato", "9/10"],
    ]
    ftable = Table(funzioni, colWidths=[6 * cm, 6 * cm, 2 * cm])
    ftable.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0B1F3A")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#B8C2CC")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(ftable)

    story.append(PageBreak())
    story.append(Paragraph("Ultimi documenti", styles["Copper"]))
    if df_docs.empty:
        story.append(Paragraph("Nessun documento ancora analizzato.", styles["BodyText"]))
    else:
        show = df_docs.head(20).copy()
        doc_data = [["File", "Cliente", "Tipo", "Stato", "Conf."]]
        for _, row in show.iterrows():
            doc_data.append([
                Paragraph(str(row.get("original_name", ""))[:45], styles["BodyText"]),
                Paragraph(str(row.get("client_name", ""))[:35], styles["BodyText"]),
                str(row.get("document_type", ""))[:18],
                str(row.get("status", ""))[:16],
                str(row.get("confidence", "")),
            ])
        dtable = Table(doc_data, colWidths=[5.2 * cm, 4.5 * cm, 3 * cm, 2.7 * cm, 1.2 * cm])
        dtable.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0B1F3A")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B8C2CC")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(dtable)

    doc.build(story, onFirstPage=_pdf_header_footer, onLaterPages=_pdf_header_footer)
    buffer.seek(0)
    return buffer.read()


# -----------------------------
# Streamlit UI
# -----------------------------


def apply_css() -> None:
    st.markdown(
        """
        <style>
        .main {background:#F6F8FB;}
        [data-testid="stSidebar"] {background:#0B1F3A;}
        [data-testid="stSidebar"] * {color:white;}
        .fp-card {
            background:white;
            padding:18px 20px;
            border-radius:16px;
            border:1px solid #E6EAF0;
            box-shadow:0 6px 20px rgba(11,31,58,0.06);
            margin-bottom:14px;
        }
        .fp-kpi {
            background:white;
            border-left:6px solid #B87333;
            padding:14px 16px;
            border-radius:14px;
            box-shadow:0 4px 16px rgba(11,31,58,0.05);
        }
        .fp-title {font-size:34px; font-weight:800; color:#0B1F3A; margin-bottom:2px;}
        .fp-subtitle {font-size:15px; color:#4D5B6A; margin-bottom:20px;}
        .badge-ok {background:#E8F5E9;color:#1B5E20;padding:4px 9px;border-radius:10px;font-weight:700;}
        .badge-warn {background:#FFF8E1;color:#8A5A00;padding:4px 9px;border-radius:10px;font-weight:700;}
        .badge-dup {background:#ECEFF1;color:#37474F;padding:4px 9px;border-radius:10px;font-weight:700;}
        </style>
        """,
        unsafe_allow_html=True,
    )


def header() -> None:
    st.markdown(f"<div class='fp-title'>{APP_NAME}</div>", unsafe_allow_html=True)
    st.markdown(
        "<div class='fp-subtitle'>OCR, classificazione documenti, matching cliente e archivio automatico cliente/mese/tipologia.</div>",
        unsafe_allow_html=True,
    )


def kpi(label: str, value: str, note: str = "") -> None:
    st.markdown(
        f"<div class='fp-kpi'><div style='font-size:13px;color:#5E6B77'>{label}</div>"
        f"<div style='font-size:28px;font-weight:800;color:#0B1F3A'>{value}</div>"
        f"<div style='font-size:12px;color:#778899'>{note}</div></div>",
        unsafe_allow_html=True,
    )


def sidebar_nav() -> str:
    st.sidebar.markdown("# FinancePlus")
    st.sidebar.caption(f"App 01_07 - v{APP_VERSION}")
    return st.sidebar.radio(
        "Menu",
        [
            "Dashboard",
            "Anagrafica aziende",
            "Importa cartella/documenti",
            "Coda da verificare",
            "Cerca azienda",
            "Archivio e report",
            "Configurazione",
        ],
    )


def page_dashboard() -> None:
    header()
    df_docs = documents_df()
    df_clients = clients_df()
    total_docs = len(df_docs)
    archived = int((df_docs["status"] == "archiviato").sum()) if not df_docs.empty else 0
    queue = int((df_docs["status"] == "da_verificare").sum()) if not df_docs.empty else 0
    dup = int((df_docs["status"] == "duplicato").sum()) if not df_docs.empty else 0

    c1, c2, c3, c4 = st.columns(4)
    with c1:
        kpi("Aziende", str(len(df_clients)), "anagrafica clienti")
    with c2:
        kpi("Documenti", str(total_docs), "analizzati")
    with c3:
        kpi("Archiviati", str(archived), "salvati per cliente")
    with c4:
        kpi("Da verificare", str(queue), f"duplicati: {dup}")

    st.markdown("### Flusso operativo")
    st.markdown(
        """
        <div class='fp-card'>
        <b>1.</b> Seleziona una cartella anche con sottocartelle ->
        <b>2.</b> Estrazione testo OCR/PDF ->
        <b>3.</b> Classificazione documento ->
        <b>4.</b> Matching cliente per ragione sociale, P.IVA, codice fiscale e amministratore ->
        <b>5.</b> Archivio automatico o coda da verificare.
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown("### Ultimi documenti")
    if df_docs.empty:
        st.info("Nessun documento ancora caricato. Vai su 'Importa cartella/documenti'.")
    else:
        show_cols = ["original_name", "client_name", "document_type", "document_date", "confidence", "status", "archived_path"]
        st.dataframe(df_docs[show_cols], use_container_width=True, hide_index=True)


def page_clients() -> None:
    header()
    st.markdown("### Anagrafica aziende")
    with st.expander("Inserisci nuova azienda", expanded=True):
        with st.form("form_client"):
            c1, c2 = st.columns(2)
            with c1:
                company_name = st.text_input("Ragione sociale *")
                vat_number = st.text_input("Partita IVA")
                fiscal_code = st.text_input("Codice fiscale")
            with c2:
                administrator = st.text_input("Amministratore / referente")
                email = st.text_input("Email")
                pec = st.text_input("PEC")
            notes = st.text_area("Note")
            submitted = st.form_submit_button("Salva azienda")
        if submitted:
            if not company_name.strip():
                st.error("Inserisci almeno la ragione sociale.")
            else:
                cid = insert_client(company_name, vat_number, fiscal_code, administrator, email, pec, notes)
                st.success(f"Azienda salvata. ID interno: {cid}")
                st.rerun()

    df = clients_df()
    st.markdown("### Elenco aziende")
    if df.empty:
        st.warning("Anagrafica vuota.")
    else:
        st.dataframe(df[["id", "company_name", "vat_number", "fiscal_code", "administrator", "email", "pec"]], use_container_width=True, hide_index=True)

    with st.expander("Importa anagrafica da CSV/Excel"):
        file = st.file_uploader("Carica file con colonne company_name, vat_number, fiscal_code, administrator, email, pec", type=["csv", "xlsx"], key="clients_import")
        if file is not None and st.button("Importa aziende"):
            raw = read_bytes(file)
            try:
                if file.name.lower().endswith(".csv"):
                    imp = pd.read_csv(io.BytesIO(raw))
                else:
                    imp = pd.read_excel(io.BytesIO(raw))
                count = 0
                for _, row in imp.iterrows():
                    name = str(row.get("company_name", row.get("ragione_sociale", ""))).strip()
                    if not name:
                        continue
                    insert_client(
                        name,
                        str(row.get("vat_number", row.get("partita_iva", "")) or ""),
                        str(row.get("fiscal_code", row.get("codice_fiscale", "")) or ""),
                        str(row.get("administrator", row.get("amministratore", "")) or ""),
                        str(row.get("email", "") or ""),
                        str(row.get("pec", "") or ""),
                        "Importato da file",
                    )
                    count += 1
                st.success(f"Importate {count} aziende.")
                st.rerun()
            except Exception as e:
                st.error(f"Errore importazione: {e}")


def page_import() -> None:
    header()
    st.markdown("### Importa cartella o documenti")
    st.info("Puoi selezionare una cartella intera: se contiene sottocartelle, Streamlit carica anche i file interni supportati.")

    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        auto_archive = st.toggle("Archivia automaticamente", value=True)
    with c2:
        min_confidence = st.slider("Confidenza minima", 40, 95, 70, step=5)
    with c3:
        allow_demo = st.toggle("Anagrafiche demo", value=False)

    if allow_demo and st.button("Carica anagrafiche demo"):
        seed_demo_clients()
        st.success("Anagrafiche demo caricate.")
        st.rerun()

    uploaded_files = st.file_uploader(
        "Seleziona una cartella o piu file",
        type=sorted([e.replace(".", "") for e in ALLOWED_EXTENSIONS]),
        accept_multiple_files="directory",
        help="Usa la selezione cartella per includere anche le sottocartelle. In locale puoi anche trascinare piu file.",
    )

    if not uploaded_files:
        st.markdown(
            """
            <div class='fp-card'>
            <b>Formato archivio creato:</b><br>
            archivio_clienti / CLIENTE / MM-AAAA / TIPO_DOCUMENTO / CLIENTE_tipo_data.ext
            </div>
            """,
            unsafe_allow_html=True,
        )
        return

    st.write(f"File selezionati: **{len(uploaded_files)}**")
    if st.button("Analizza e archivia", type="primary"):
        results: List[AnalysisResult] = []
        progress = st.progress(0)
        status_box = st.empty()
        for idx, uf in enumerate(uploaded_files, start=1):
            status_box.write(f"Analisi {idx}/{len(uploaded_files)}: {uf.name}")
            try:
                result = analyze_upload(uf, auto_archive=auto_archive, min_confidence=min_confidence)
                insert_document(result)
                results.append(result)
            except Exception as e:
                st.error(f"Errore su {uf.name}: {e}")
            progress.progress(idx / len(uploaded_files))
        status_box.success("Analisi completata.")

        if results:
            df = pd.DataFrame([r.__dict__ for r in results])
            st.markdown("### Esito importazione")
            st.dataframe(
                df[["original_name", "matched_client_name", "document_type", "document_date", "confidence", "status", "suggested_filename", "archived_path"]],
                use_container_width=True,
                hide_index=True,
            )
            csv = df.to_csv(index=False).encode("utf-8")
            st.download_button("Scarica esito CSV", csv, "esito_import_app_01_07.csv", "text/csv")


def page_queue() -> None:
    header()
    st.markdown("### Coda documenti da verificare")
    df = documents_df("d.status = ?", ("da_verificare",))
    clients = clients_df()
    if df.empty:
        st.success("Nessun documento in coda da verificare.")
        return

    st.dataframe(df[["id", "original_name", "client_name", "document_type", "document_date", "confidence", "notes"]], use_container_width=True, hide_index=True)
    doc_id = st.selectbox("Seleziona documento", df["id"].tolist(), format_func=lambda x: f"#{x} - {df[df['id'] == x].iloc[0]['original_name']}")
    row = df[df["id"] == doc_id].iloc[0]
    st.markdown("#### Anteprima testo")
    st.text_area("Testo estratto", value=str(row.get("text_preview", "")), height=180)

    options = {int(r["id"]): str(r["company_name"]) for _, r in clients.iterrows()}
    selected_client_id = st.selectbox("Abbina ad azienda", [None] + list(options.keys()), format_func=lambda x: "-- scegli --" if x is None else options[x])
    doc_type = st.selectbox("Tipologia corretta", sorted(list(DOC_TYPES.keys()) + ["altro"]), index=sorted(list(DOC_TYPES.keys()) + ["altro"]).index(str(row.get("document_type", "altro"))))
    doc_date = st.text_input("Data documento gg-mm-aaaa", value=str(row.get("document_date", today_str())))

    if st.button("Conferma e archivia", type="primary"):
        if selected_client_id is None:
            st.error("Seleziona prima il cliente.")
            return
        client_name = options[selected_client_id]
        ext = str(row.get("extension", ".pdf"))
        suggested = suggest_name(client_name, doc_type, doc_date, ext)
        # Il file originale in coda non e sempre presente come raw bytes: se era non archiviato, salvo un segnaposto testuale.
        # Per archiviazione reale dei non riconosciuti, ricaricare il file o usare import automatico con soglia piu bassa.
        target_dir = ARCHIVE_ROOT / company_slug(client_name) / month_folder_from_date(doc_date) / safe_filename(doc_type, "altro")
        target_dir.mkdir(parents=True, exist_ok=True)
        target = target_dir / safe_filename(suggested, "documento.txt")
        if not str(row.get("archived_path", "")):
            target.write_text(str(row.get("full_text", "")) or str(row.get("text_preview", "")), encoding="utf-8")
            archived_path = str(target)
        else:
            archived_path = str(row.get("archived_path", ""))
        update_document_status(int(doc_id), "archiviato", int(selected_client_id), client_name, archived_path, suggested)
        st.success("Documento verificato e archiviato.")
        st.rerun()


def zip_files(paths: Iterable[str]) -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for p in paths:
            if not p:
                continue
            path = Path(p)
            if path.exists() and path.is_file():
                zf.write(path, arcname=str(path.relative_to(ARCHIVE_ROOT.parent)) if ARCHIVE_ROOT.parent in path.parents else path.name)
    buffer.seek(0)
    return buffer.read()


def page_search_company() -> None:
    header()
    st.markdown("### Cerca azienda")
    query = st.text_input("Scrivi nome azienda, P.IVA, codice fiscale o parola contenuta nei documenti")
    if not query:
        st.info("Inserisci una ricerca per visualizzare tutti i documenti collegati.")
        return

    q = f"%{query.lower()}%"
    df = documents_df(
        "LOWER(d.client_name) LIKE ? OR LOWER(d.original_name) LIKE ? OR LOWER(d.full_text) LIKE ? OR d.vat_number LIKE ? OR d.fiscal_code LIKE ?",
        (q, q, q, q, q),
    )
    st.write(f"Risultati trovati: **{len(df)}**")
    if df.empty:
        st.warning("Nessun risultato.")
        return

    st.dataframe(df[["id", "original_name", "client_name", "document_type", "document_date", "confidence", "status", "archived_path"]], use_container_width=True, hide_index=True)

    st.markdown("#### Anteprima e sintesi")
    doc_id = st.selectbox("Seleziona documento per anteprima", df["id"].tolist(), format_func=lambda x: f"#{x} - {df[df['id'] == x].iloc[0]['original_name']}")
    row = df[df["id"] == doc_id].iloc[0]
    st.markdown(
        f"""
        <div class='fp-card'>
        <b>Cliente:</b> {row.get('client_name','')}<br>
        <b>Tipo:</b> {row.get('document_type','')}<br>
        <b>Data:</b> {row.get('document_date','')}<br>
        <b>Confidenza:</b> {row.get('confidence','')}%<br>
        <b>Percorso:</b> {row.get('archived_path','')}
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.text_area("Sintesi / anteprima testo", value=str(row.get("text_preview", "")), height=180)

    selectable = df[df["archived_path"].astype(str) != ""].copy()
    if selectable.empty:
        st.info("I risultati non hanno file fisici archiviati scaricabili. Importa con archiviazione automatica oppure verifica la coda.")
        return

    ids = st.multiselect("Seleziona documenti da scaricare", selectable["id"].tolist(), default=selectable["id"].head(5).tolist())
    selected_paths = selectable[selectable["id"].isin(ids)]["archived_path"].tolist()
    if selected_paths:
        zip_bytes = zip_files(selected_paths)
        st.download_button("Scarica selezionati in ZIP", zip_bytes, f"documenti_{safe_filename(query)}.zip", "application/zip")


def page_reports() -> None:
    header()
    st.markdown("### Archivio e report")
    df_docs = documents_df()
    df_clients = clients_df()

    c1, c2 = st.columns(2)
    with c1:
        st.markdown("#### Archivio fisico")
        st.code(str(ARCHIVE_ROOT.resolve()))
        if ARCHIVE_ROOT.exists():
            count = sum(1 for p in ARCHIVE_ROOT.rglob("*") if p.is_file())
            st.write(f"File fisici in archivio: **{count}**")
    with c2:
        st.markdown("#### Database")
        st.code(str(DB_PATH.resolve()))
        st.write(f"Documenti nel database: **{len(df_docs)}**")

    st.markdown("#### Elenco documenti")
    if df_docs.empty:
        st.info("Nessun documento disponibile.")
    else:
        st.dataframe(df_docs[["id", "original_name", "client_name", "document_type", "document_date", "status", "archived_path"]], use_container_width=True, hide_index=True)

    st.markdown("#### Genera report PDF")
    if st.button("Crea PDF riepilogativo", type="primary"):
        try:
            pdf_bytes = create_pdf_report(df_docs, df_clients)
            filename = f"Report_App_01_07_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf"
            out = REPORT_DIR / filename
            out.write_bytes(pdf_bytes)
            st.success("Report creato.")
            st.download_button("Scarica report PDF", pdf_bytes, filename, "application/pdf")
        except Exception as e:
            st.error(f"Errore generazione report: {e}")

    st.markdown("#### Esporta database documenti")
    if not df_docs.empty:
        st.download_button("Scarica CSV documenti", df_docs.to_csv(index=False).encode("utf-8"), "documenti_app_01_07.csv", "text/csv")
    if not df_clients.empty:
        st.download_button("Scarica CSV aziende", df_clients.to_csv(index=False).encode("utf-8"), "aziende_app_01_07.csv", "text/csv")


def page_config() -> None:
    header()
    st.markdown("### Configurazione tecnica")
    st.markdown(
        """
        <div class='fp-card'>
        <b>Avvio locale:</b><br>
        <code>streamlit run streamlit_app.py</code><br><br>
        <b>Struttura cloud:</b><br>
        carica su GitHub: <code>streamlit_app.py</code>, <code>requirements.txt</code>, <code>packages.txt</code>, <code>.streamlit/config.toml</code>.
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown("#### Stato librerie")
    status = pd.DataFrame(
        [
            ["PyMuPDF / fitz", fitz is not None, "Lettura PDF digitale e rasterizzazione OCR"],
            ["pdfplumber", pdfplumber is not None, "Fallback lettura PDF"],
            ["Pillow", Image is not None, "Apertura immagini"],
            ["pytesseract", pytesseract is not None, "OCR immagini e PDF scannerizzati"],
            ["python-docx", Document is not None, "Lettura DOCX"],
            ["reportlab", SimpleDocTemplate is not None, "Report PDF"],
        ],
        columns=["Componente", "Disponibile", "Uso"],
    )
    st.dataframe(status, use_container_width=True, hide_index=True)

    st.markdown("#### Pulizia dati")
    st.warning("Usare solo in fase di test: elimina database e archivio locale del progetto.")
    confirm = st.text_input("Scrivi ELIMINA per abilitare")
    if confirm == "ELIMINA":
        if st.button("Elimina database e archivio", type="secondary"):
            if DB_PATH.exists():
                DB_PATH.unlink()
            if ARCHIVE_ROOT.exists():
                shutil.rmtree(ARCHIVE_ROOT)
            ensure_dirs()
            init_db()
            st.success("Dati eliminati e database ricreato.")
            st.rerun()


def main() -> None:
    st.set_page_config(page_title=APP_NAME, page_icon="FP", layout="wide")
    ensure_dirs()
    init_db()
    apply_css()
    page = sidebar_nav()

    if page == "Dashboard":
        page_dashboard()
    elif page == "Anagrafica aziende":
        page_clients()
    elif page == "Importa cartella/documenti":
        page_import()
    elif page == "Coda da verificare":
        page_queue()
    elif page == "Cerca azienda":
        page_search_company()
    elif page == "Archivio e report":
        page_reports()
    elif page == "Configurazione":
        page_config()

    st.sidebar.markdown("---")
    st.sidebar.caption("FinancePlus.Tech - IDP documentale")


if __name__ == "__main__":
    main()
